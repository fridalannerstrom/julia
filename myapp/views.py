import os
import io
import re
import json
import textwrap
import openpyxl
import markdown2
from markdown2 import markdown
import math
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.utils.safestring import mark_safe
from dotenv import load_dotenv
from openai import OpenAI
from django.conf import settings
from .models import Prompt, ChatSession, ChatMessage, ChatAttachment
from django.http import StreamingHttpResponse, JsonResponse
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, redirect
from django.views.decorators.http import require_POST
from bs4 import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import base64
from docx.shared import Inches
from io import BytesIO
from .models import Report
from django.utils import timezone
from django.urls import reverse

# ──────────────────────────────────────────────────────────────────────────────
# Miljö
# ──────────────────────────────────────────────────────────────────────────────


load_dotenv()
if os.path.exists("env.py"):
    import env  # noqa: F401

client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
    timeout=20,      # global timeout per request
    max_retries=2,   # här är den OK
)

REPORT_CONTEXT_KEYS = [
    # styrning
    "step",
    "ratings_json",

    # inputs / rådata
    "test_text",
    "intervju_text",
    "cv_text",
    "uploaded_files_markdown",
    "job_ad_text",
    "motivation_notes",
    "logical_score",
    "verbal_score",
    "candidate_first_name",
    "candidate_last_name",
    "candidate_name",
    "candidate_role",
    "selected_motivation_keys",

    # AI-texter
    "tq_fardighet_text",
    "tq_motivation_text",
    "leda_text",
    "mod_text",
    "sjalkannedom_text",
    "strategi_text",
    "kommunikation_text",
    "sur_text",
    "slutsats_text",
]


def _report_title_from_context(ctx: dict) -> str:
    name = (ctx.get("candidate_name") or "").strip()
    role = (ctx.get("candidate_role") or "").strip()
    if name and role:
        return f"{name} – {role}"
    return name or role or "Ny rapport"

def _extract_report_data_from_context(ctx: dict) -> dict:
    data = {}
    for k in REPORT_CONTEXT_KEYS:
        if k in ctx:
            data[k] = ctx.get(k)
    return data

def _apply_report_data_to_context(ctx: dict, data: dict) -> dict:
    for k, v in (data or {}).items():
        # vi vill inte att random keys ska skräpa ner, men ok att tillåta våra
        if k in REPORT_CONTEXT_KEYS:
            ctx[k] = v
    return ctx

def _get_report_or_404(report_id):
    return get_object_or_404(Report, id=report_id, deleted_at__isnull=True)

def _ensure_report(request, ctx: dict):
    """
    Skapar eller hämtar report baserat på report_id (GET/POST).
    """
    report_id = request.POST.get("report_id") or request.GET.get("report_id")
    if report_id:
        rep = _get_report_or_404(report_id)
        return rep

    # Ingen report_id => skapa ny rapport vid start (GET step=1)
    rep = Report.objects.create(
        created_by=request.user,
        current_step=int(ctx.get("step") or 1),
        title=_report_title_from_context(ctx),
        data=_extract_report_data_from_context(ctx),
    )
    return rep

def _save_report_state(rep: Report, ctx: dict):
    rep.current_step = int(ctx.get("step") or rep.current_step or 1)
    rep.title = _report_title_from_context(ctx)
    rep.data = _extract_report_data_from_context(ctx)
    rep.save(update_fields=["current_step", "title", "data", "updated_at"])


# ── NYTT: gemensamma rubriknycklar i rätt ordning ────────────────────────────
SECTION_KEYS = [
    ("tq_fardighet_text", "TQ Färdighet"),
    ("tq_motivation_text", "TQ Motivation"),
    ("leda_text", "Leda, utveckla och engagera"),
    ("mod_text", "Mod och handlingskraft"),
    ("sjalkannedom_text", "Självkännedom och emotionell stabilitet"),
    ("strategi_text", "Strategiskt tänkande och anpassningsförmåga"),
    ("kommunikation_text", "Kommunikation och samarbete"),
]

# ── Koppling från STIVE-kompetenser -> (sektion_key, svensk_rad) ────────────
HEADER_TO_TARGET = {
    # Leda, utveckla och engagera
    "directing others":        ("leda_utveckla_och_engagera", "Leda andra"),
    "engaging others":         ("leda_utveckla_och_engagera", "Engagera andra"),
    "delegating":              ("leda_utveckla_och_engagera", "Delegera"),
    "developing others":       ("leda_utveckla_och_engagera", "Utveckla andra"),

    # Mod och handlingskraft
    "decisiveness":            ("mod_och_handlingskraft", "Beslutsamhet"),
    "integrity":               ("mod_och_handlingskraft", "Integritet"),
    "managing conflict":       ("mod_och_handlingskraft", "Hantera konflikter"),

    # Självkännedom och emotionell stabilitet
    "self-awareness":          ("sjalkannedom_och_emotionell_stabilitet", "Självmedvetenhet"),
    "resilience":              ("sjalkannedom_och_emotionell_stabilitet", "Uthållighet"),

    # Strategiskt tänkande och anpassningsförmåga
    "strategic focus":         ("strategiskt_tankande_och_anpassningsformaga", "Strategiskt fokus"),
    "adaptability":            ("strategiskt_tankande_och_anpassningsformaga", "Anpassningsförmåga"),

    # Kommunikation och samarbete
    "teamwork":                ("kommunikation_och_samarbete", "Teamarbete"),
    "influencing":             ("kommunikation_och_samarbete", "Inflytelserik"),
}

MOTIVATION_FACTORS = {
    "anknytning": {
        "label": "Anknytning",
        "definition": (
            "Värdesätter social interaktion med andra på jobbet, detta är en primär drivkraft. "
            "Trivs med känslan av att tillhöra ett team och drivs av att utveckla starka band "
            "och goda arbetsrelationer."
        ),
    },
    "kundservice": {
        "label": "Kundservice",
        "definition": (
            "Trivs i en kundorienterad roll där de får hjälpa och stödja andra. Lägger värde på "
            "att bygga relationer med kunder och förstå kundernas behov, och lägger stor vikt vid "
            "att möta och överträffa kunders behov och förväntningar."
        ),
    },
    "work_life_balance": {
        "label": "Work life balance",
        "definition": (
            "Balans mellan arbete och privatliv, så de kan uppfylla sina åtaganden både i arbetet "
            "och utanför. Söker tydliga gränser så de kan stänga av när arbetsdagen är slut och "
            "ägna tid åt andra aktiviteter som är viktiga."
        ),
    },
    "manniskors_utveckling": {
        "label": "Människors utveckling",
        "definition": (
            "Hjälpa och stödja andra och aktivt bidra till deras utveckling. Trivs med att vägleda "
            "andra genom svåra tider och hjälper andra att övervinna motgångar eller utmaningar, "
            "samt erbjuder proaktivt stöd så att andra kan växa."
        ),
    },
    "stabilitet": {
        "label": "Stabilitet",
        "definition": (
            "Känna trygghet på jobbet och föredra att arbeta i en bransch eller organisation som "
            "upplevs stabil. Lägger vikt vid att känna att deras position är säker och värdesätter "
            "stabilitet i sin roll."
        ),
    },
    "auktoritet": {
        "label": "Auktoritet",
        "definition": (
            "Värdesätter möjligheten att ta på sig ledande befattningar och möjlighet till befordran. "
            "Uppskattar roller där de kan ange riktningen, organisera andra och säkerställa att arbetet "
            "utförs som det ska."
        ),
    },
    "forvarv": {
        "label": "Förvärv",
        "definition": (
            "Motiveras av ekonomisk belöning och att arbeta i en välbetald roll så de har råd med de "
            "saker de vill ha utan att behöva oroa sig för kostnader. Fokus på lön och förmåner är "
            "starkt, och attraktiv ersättning är motiverande."
        ),
    },
    "autonomi": {
        "label": "Autonomi",
        "definition": (
            "Friheten att fatta egna beslut om hur de ska planera sin tid och organisera sitt arbete. "
            "Presterar som bäst när de får förtroende att sätta sin egen arbetsriktning och slutföra "
            "arbetet på det sätt de tycker är lämpligt."
        ),
    },
    "erkannande": {
        "label": "Erkännande",
        "definition": (
            "Motiveras av erkännande, beröm och uppskattning. Prioriterar aktiviteter eller projekt "
            "som är synliga i organisationen så att andra kan se arbetet de gör och värdet de tillför. "
            "Positiv feedback och beröm är starkt motiverande."
        ),
    },
    "gora_skillnad": {
        "label": "Göra skillnad",
        "definition": (
            "Känna att de gör en positiv skillnad för andra. Vill vara del av en organisation vars mål "
            "och syfte upplevs som ’värdigt’ och uppskattar att se hur deras arbete har en gynnsam "
            "inverkan på samhället i stort."
        ),
    },
    "prestation": {
        "label": "Prestation",
        "definition": (
            "Får energi av definierade och utmanande mål att arbeta mot. Möjligheter som utmanar "
            "professionellt har stor betydelse, och de trivs med konkurrens och en miljö där de kan "
            "jämföra sin prestation med andras."
        ),
    },
    "kvalitet": {
        "label": "Kvalitet",
        "definition": (
            "Vill leverera högkvalitativa resultat på en konsekvent basis. Är stolta över att producera "
            "arbete som ligger i linje med både vad de lovat och vad som förväntas av dem, och de "
            "säkerställer att arbetet slutförs så perfekt som möjligt."
        ),
    },
    "larande": {
        "label": "Lärande",
        "definition": (
            "Drivs av möjligheten att lära sig och utveckla sin expertis. Trivs i en miljö där det finns "
            "gott om möjligheter till fortlöpande utbildning och utveckling, och motiveras av att bemästra "
            "nya färdigheter."
        ),
    },
    "etik": {
        "label": "Etik",
        "definition": (
            "Upprätthåller hög etisk standard för sig själv och sin organisation. Följer branschriktlinjer "
            "och krav på bästa praxis. Är stolta över sin professionella integritet och trivs i en organisation "
            "vars etiska värderingar ligger i linje med de egna."
        ),
    },
    "marknadsvarde": {
        "label": "Marknadsvärde",
        "definition": (
            "Trivs i en roll som är kommersiellt fokuserad och bidrar till organisationens ekonomiska "
            "framgång. Motiveras av att se kopplingen mellan sin roll och organisationens resultat och "
            "får energi av att se hur deras arbete leder till kommersiell vinst."
        ),
    },
    "nyfikenhet": {
        "label": "Nyfikenhet",
        "definition": (
            "Vill utforska och upptäcka nya saker. Eftersträvar kunskap och erfarenhet, tar gärna möjligheter "
            "till nya upplevelser, ny information och att utforska sådant som fångar deras intresse."
        ),
    },
    "kreativitet": {
        "label": "Kreativitet",
        "definition": (
            "Motiveras av att utöva sin kreativitet, tänka utanför ramarna, generera nya idéer och identifiera "
            "nya lösningar. Deras idéer är ofta användbara, inte bara originella för sakens skull."
        ),
    },
    "gladje": {
        "label": "Glädje",
        "definition": (
            "Värdesätter en miljö där man trivs med sitt arbete och kan ha roligt med kollegor. Trivs i en "
            "lättsam miljö där de kan uttrycka sin lekfulla sida, dela humor med kollegor och bli accepterade "
            "för den de är."
        ),
    },
    "variation": {
        "label": "Variation",
        "definition": (
            "Motiveras av en roll med mycket variation och möjlighet att arbeta med olika saker. Trivs i en "
            "mångsidig roll med en blandning av uppgifter och projekt där de kan rikta sin uppmärksamhet mot "
            "olika arbetskrav och använda sina olika färdigheter."
        ),
    },
    "risk": {
        "label": "Risk",
        "definition": (
            "Motiveras av risk i arbetet och att få möta situationer där utfallet är osäkert. Är bekväma med att "
            "ta riskfyllda beslut och trivs i roller där de får uppleva spänningen i att ta en risk och se om "
            "den lönar sig."
        ),
    },
}

def _to_str(v):
    if v is None:
        return ""
    if isinstance(v, (list, dict)):
        return json.dumps(v, ensure_ascii=False, indent=2)
    return str(v)


# ── NYTT: liten wrapper för OpenAI-anrop per rubrik ───────────────────────────
def _run_openai(prompt_text: str, style: str, **vars_) -> str:
    def _to_str(v):
        """
        Säker konvertering till str så .replace() aldrig får list/dict.
        - list/dict -> JSON (lätt för AI att läsa)
        - None -> ""
        - annat -> str(...)
        """
        if v is None:
            return ""
        if isinstance(v, (list, dict)):
            try:
                return json.dumps(v, ensure_ascii=False, indent=2)
            except Exception:
                return str(v)
        return str(v)

    try:
        # ✅ DEBUG: se om motivation_notes finns och hur lång den är
        mn = vars_.get("motivation_notes", None)
        print("DEBUG motivation_notes type:", type(mn))
        print("DEBUG motivation_notes len:", len(mn) if isinstance(mn, str) else "not str")
        print("DEBUG motivation_notes preview:", repr((mn or "")[:200]))

        # ✅ DEBUG: se om placeholdern ens finns i prompten
        print("DEBUG prompt has {motivation_notes}:", "{motivation_notes}" in prompt_text)

        # Gör en kopia av prompt_text så vi inte muterar original
        pt = str(prompt_text or "")

        # Behåll dina två "snabba" replacements (men säkra)
        pt = pt.replace("{excel_text}", _to_str(vars_.get("excel_text", "")))
        pt = pt.replace("{intervju_text}", _to_str(vars_.get("intervju_text", "")))

        # ✅ stöd för fler placeholders utan att krascha (fixen!)
        for k, v in vars_.items():
            placeholder = "{" + k + "}"
            pt = pt.replace(placeholder, _to_str(v))

        filled = (style or "") + "\n\n" + pt

        # ✅ DEBUG: kolla om motivation_notes faktiskt hamnade i filled
        mn_preview = (mn or "")
        idx = filled.find(str(mn_preview)[:50]) if isinstance(mn_preview, str) else -1
        print("DEBUG motivation_notes appears in filled:", idx != -1)
        print("DEBUG filled length:", len(filled))

        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": filled}],
            temperature=0.2,
            max_tokens=900,
            timeout=20,
        )
        content = resp.choices[0].message.content
        if not content:
            raise ValueError("Tomt svar från OpenAI")
        return content.strip()

    except Exception as e:
        print("⚠️ OpenAI error in _run_openai:", repr(e))
        return (
            "Tyvärr tog AI-svaret för lång tid eller gick inte att hämta just nu. "
            "Försök igen om en liten stund."
        )


def _round_to_1_5(x) -> int:
    """
    Runda till heltal mellan 1-5.
    Ex:
      2.2 -> 2
      2.7 -> 3
    (klassisk .5 uppåt, inte bankers rounding)
    """
    try:
        v = float(str(x).replace(",", "."))
    except Exception:
        return 3

    if v < 1:
        v = 1
    if v > 5:
        v = 5

    return int(math.floor(v + 0.5))

def _build_sidebar_context(owner, step, context, ratings_json_str):
    """
    Bygger ett context-paket till sidobarschattens AI.
    Innehåller kandidatinfo, test/intervju/CV + aktuell sektion och dess prompt(er).
    """
    base = {
        "step": step,
        "candidate_name": context.get("candidate_name", ""),
        "candidate_role": context.get("candidate_role", ""),
        "test_text": context.get("test_text", ""),
        "intervju_text": context.get("intervju_text", ""),
        "cv_text": context.get("uploaded_files_markdown", "") or context.get("cv_text", ""),
        "ratings_json": ratings_json_str or context.get("ratings_json", ""),
        "sections": [],
    }

    def add_section(field_label, field_key, prompt_name):
        try:
            prompt_obj = Prompt.objects.get(user=owner, name=prompt_name)
            prompt_text = prompt_obj.text
        except Prompt.DoesNotExist:
            prompt_text = ""

        base["sections"].append({
            "field_label": field_label,         # t.ex. "Leda, utveckla och engagera"
            "field_key": field_key,             # t.ex. "leda_text"
            "current_text": context.get(field_key, ""),
            "prompt_name": prompt_name,         # t.ex. "leda"
            "prompt_text": prompt_text,
        })

    # Vilka sektioner hör till vilket steg?
    if step == 2:
        # Steg 2 – TQ Färdighet
        add_section("TQ Färdighet", "tq_fardighet_text", "tq_fardighet")

    elif step == 3:
        # Steg 3 – TQ Motivation
        add_section("TQ Motivation", "tq_motivation_text", "tq_motivation")

    elif step == 4:
        add_section("Leda, utveckla och engagera", "leda_text", "leda")

    elif step == 5:
        add_section("Mod och handlingskraft", "mod_text", "mod")

    elif step == 6:
        add_section(
            "Självkännedom och emotionell stabilitet",
            "sjalkannedom_text",
            "sjalkannedom",
        )

    elif step == 7:
        add_section(
            "Strategiskt tänkande och anpassningsförmåga",
            "strategi_text",
            "strategi",
        )

    elif step == 8:
        add_section("Kommunikation och samarbete", "kommunikation_text", "kommunikation")

    elif step == 9:
        add_section(
            "Styrkor / Utvecklingsområden / Riskbeteenden",
            "sur_text",
            "styrkor_utveckling_risk",
        )

    elif step == 10:
        add_section(
            "Sammanfattande slutsats",
            "slutsats_text",
            "sammanfattande_slutsats",
        )

    return base

def _ai_text_and_ratings_for_section(config, base_prompt_text, style, excel_text, intervju_text=""):
    """
    Kör en prompt som:
      - först skriver en kort textbedömning
      - sedan ger RATINGS_JSON med 1–5 för rätt subskalor för denna sektion.
    Returnerar (full_text, ratings_dict, debug_str).
    """
    section_key = config["section_key"]
    subscales = config["subscales"]

    # Bygg tydlig instruktion som läggs till befintlig prompttext
    rating_instr_lines = [
        "",
        "Viktigt:",
        "1. Baserat på testdatan (och intervju om finns), bedöm varje delkompetens på en skala 1–5:",
        "   1 = Utrymme för utveckling",
        "   2 = Tillräcklig",
        "   3 = God",
        "   4 = Mycket god",
        "   5 = Utmärkt",
        "2. Säkerställ att texten du skriver stämmer överens med de poäng du sätter.",
        "3. Avsluta ALLTID ditt svar med exakt följande format:",
        "   ### RATINGS_JSON",
        "   {",
        f'     "{section_key}": {{'
    ]
    rating_instr_lines += [
        f'       "{sub}": <heltal 1-5>,'
        for sub in subscales[:-1]
    ]
    rating_instr_lines.append(
        f'       "{subscales[-1]}": <heltal 1-5>'
    )
    rating_instr_lines.append("     }")
    rating_instr_lines.append("   }")
    rating_instr_lines.append("")
    rating_instr = "\n".join(rating_instr_lines)

    full_prompt = base_prompt_text.strip() + "\n\n" + rating_instr

    # Kör mot OpenAI
    full_text = _run_openai(
        full_prompt,
        style,
        excel_text=_trim(excel_text),
        intervju_text=_trim(intervju_text or ""),
    )

    # Försök plocka JSON
    parsed = _safe_json_from_text(full_text) or {}
    sec_ratings = parsed.get(section_key, {})

    # Fyll in defaults om något saknas
    cleaned = {}
    debug_lines = []
    for sub in subscales:
        raw = sec_ratings.get(sub)
        try:
            v = int(raw)
        except Exception:
            v = 3
            debug_lines.append(
                f"{section_key}/{sub}: saknade eller ogiltig poäng ('{raw}'), satt till 3"
            )
        v = max(1, min(5, v))
        cleaned[sub] = v
        debug_lines.append(f"{section_key}/{sub}: {v}")

    return full_text, cleaned, "\n".join(debug_lines)

from django.contrib.auth import get_user_model
from django.conf import settings

def get_prompt_owner(fallback_user=None):
    """
    Returnerar den användare vars prompter är 'globala' för hela systemet.
    Typiskt: Veronika. Om inte hittad -> fallback_user.
    """
    User = get_user_model()
    username = getattr(settings, "PROMPT_OWNER_USERNAME", None)

    if username:
        try:
            return User.objects.get(username=username)
        except User.DoesNotExist:
            pass

    # Fallback om något blivit fel
    return fallback_user

# ── NYTT: gör _ratings_table_html konfigurerbar (header av/på) ───────────────
def _ratings_table_html(
    ratings: dict,
    section_filter=None,
    include_css: bool = True
) -> str:
    # Ordning på sektionerna
    default_order = [
        ("leda_utveckla_och_engagera", "Leda, utveckla och engagera"),
        ("mod_och_handlingskraft", "Mod och handlingskraft"),
        ("sjalkannedom_och_emotionell_stabilitet", "Självkännedom och emotionell stabilitet"),
        ("strategiskt_tankande_och_anpassningsformaga", "Strategiskt tänkande och anpassningsförmåga"),
        ("kommunikation_och_samarbete", "Kommunikation och samarbete"),
    ]
    section_order = section_filter or default_order

    def row(section_key: str, label: str, val: int):
        # Hämta ev. definitionstext
        desc = SUBSCALE_DESCRIPTIONS.get(section_key, {}).get(label, "")
        desc_html = f'<div class="dn-sub-desc">{desc}</div>' if desc else ""

        # 5 cirklar, en fylld
        cells = []
        for i in range(1, 6):
            active_class = " dn-dot--active" if val == i else ""
            cells.append(
                f'<td class="dn-cell"><span class="dn-dot{active_class}"></span></td>'
            )

        return (
            '<tr>'
            '<th class="dn-sub">'
            f'<div class="dn-sub-title ">{label}</div>'
            f'<div class="dn-sub-desc ">{desc_html}</div>'
            '</th>'
            f'{"".join(cells)}'
            '</tr>'
        )

    sections_html = []

    for key, title in section_order:
        if key not in ratings:
            continue

        section_ratings = ratings.get(key, {})
        rows_html = []

        target_labels = TARGETS.get(key)
        if target_labels:
            for label in target_labels:
                raw_score = section_ratings.get(label, 3)
                try:
                    v = int(raw_score)
                except Exception:
                    v = 3
                v = max(1, min(5, v))
                rows_html.append(row(key, label, v))
        else:
            for label, raw_score in section_ratings.items():
                try:
                    v = int(raw_score)
                except Exception:
                    v = 3
                v = max(1, min(5, v))
                rows_html.append(row(key, label, v))

        if rows_html:
            sections_html.append(f"""
            <div class="dn-section">
              <table class="dn-table">
                <tbody>
                  {''.join(rows_html)}
                </tbody>
              </table>
            </div>
            """)

        css = """
        <style>
        .dn-section{
            margin:24px 0;
        }

        .dn-table{
            width:100%;
            border-collapse:collapse;
        }

        .dn-table tr + tr{
            border-top:1px solid #e5e7eb;
        }

        .dn-sub {
            font-weight:500;
            padding:10px 10px;
            white-space:normal;
            color:#111827;
            font-family: 'Segoe UI', Calibri, sans-serif;
            font-size:1.2rem;          /* större grundtext i vänsterkolumnen */
        }

        .dn-sub-title {
            font-weight:600;
            font-size:1.4rem;          /* större rubrik */
            margin-bottom:4px;
        }

        .dn-sub-desc {
            font-weight:400;
            font-size:1.1rem;           /* större brödtext */
            color:#4b5563;
            line-height:1.5;
        }

        .dn-cell{
            text-align:center;
            padding:12px 4px;           /* lite mer vertikal luft runt prickarna */
        }

        .dn-dot{
            display:inline-flex;
            width:25px;
            height:25px;
            border-radius:999px;
            border:1px solid #BEBEBE;
            background:#E4E4E4;
            box-shadow:0 0 0 1px rgba(255,255,255,0.9);
        }

        .dn-dot--active{
            background:#7b2cbf;
            border-color:#7b2cbf;
            box-shadow:0 0 0 3px rgba(123,44,191,0.18);
        }

        /* ─────────────────────────────
        EXPORT-VY (html2canvas)
        ───────────────────────────── */

        .rating-export{
            width:900px;          /* smalare = högre bild i Word */
            max-width:900px;
            margin:0 auto;
            background:#ffffff;
            padding:2px 2px;
        }

        .rating-export .dn-table{
            width:100%;
            table-layout:fixed;
        }

        /* vänster textkolumn */
        .rating-export .dn-sub{
            width:360px;
            max-width:360px;
        }

        /* prick-cellerna får dela på resten */
        .rating-export .dn-cell{
            width:auto;
        }
        </style>

        """

    return (css if include_css else "") + "\n".join(sections_html)


# --- Målmönster: vad i Excel-raden betyder vilken skala-rad? -----------------
# Flera varianter/engelska namn om dina mallar ändras.
TARGET_PATTERNS = {
    "leda_utveckla_och_engagera": {
        "Leda andra":        [r"leda\s+andra", r"leading\s+others"],
        "Engagera andra":    [r"engagera\s+andra", r"engag(e|era)\w*\s+others"],
        "Delegera":          [r"delegera", r"delegat\w*"],
        "Utveckla andra":    [r"utveckla\s+andra", r"develops?\s+others"],
    },
    "mod_och_handlingskraft": {
        "Beslutsamhet":      [r"beslutsamhet", r"decisiv\w*"],
        "Integritet":        [r"integritet", r"integrit(y|et)"],
        "Hantera konflikter":[r"hantera\s+konflikter", r"conflict\s+(management|handling)"],
    },
    "sjalkannedom_och_emotionell_stabilitet": {
        "Självmedvetenhet":  [r"självmedvet(enhet)?", r"self[-\s]?awareness"],
        "Uthållighet":       [r"uthållighet", r"resilien(ce|s)", r"perseverance"],
    },
    "strategiskt_tankande_och_anpassningsformaga": {
        "Strategiskt fokus": [r"strateg(iskt|ic)\s+(fokus|focus|thinking)"],
        "Anpassningsförmåga":[r"anpassningsf(ö|o)rm(å|a)ga", r"adaptab\w*", r"adapting\s+to\s+change"],
    },
    "kommunikation_och_samarbete": {
        "Teamarbete":        [r"teamarbete", r"team\s*work"],
        "Inflytelserik":     [r"inflytelserik", r"influenc\w*", r"persuasive"],
    },
}


TARGETS = {
    "leda_utveckla_och_engagera": [
        "Leda andra",
        "Engagera andra",
        "Delegera",
        "Utveckla andra",
    ],
    "mod_och_handlingskraft": [
        "Beslutsamhet",
        "Integritet",
        "Hantera konflikter",
    ],
    "sjalkannedom_och_emotionell_stabilitet": [
        "Självmedvetenhet",
        "Uthållighet",
    ],
    "strategiskt_tankande_och_anpassningsformaga": [
        "Strategiskt fokus",
        "Anpassningsförmåga",
    ],
    "kommunikation_och_samarbete": [
        "Teamarbete",
        "Inflytelserik",
    ],
}

SUBSCALE_DESCRIPTIONS = {
    "leda_utveckla_och_engagera": {
        "Leda andra": (
            "Ger tydlig riktning och följer upp så att mål och uppdrag uppnås."
        ),
        "Engagera andra": (
            "Skapar engagemang och delaktighet genom att visa intresse, energi och närvaro."
        ),
        "Delegera": (
            "Fördelar uppgifter utifrån kompetens och tillgänglighet och följer upp utan att detaljstyra."
        ),
        "Utveckla andra": (
            "Ger återkoppling och skapar möjligheter till lärande för att stärka andras utveckling."
        ),
    },
    "mod_och_handlingskraft": {
        "Beslutsamhet": (
            "Fattar bra beslut i rätt tid på sund logik och vettiga resonemang, handlar med övertygelse när ett beslut behöver tas, även med begränsad information tillhanda. "
        ),
        "Integritet": (
            "Visar prov på höga etiska standarder och arbetar på ett autentiskt och ärligt sätt, svarar på etiska konflikter med integritet. "
        ),
        "Hantera konflikter": (
            "Hanterar och löser konflikter och oenigheter bland andra, taktfullt men ändå öppet. "
        ),
    },
    "sjalkannedom_och_emotionell_stabilitet": {
        "Självmedvetenhet": (
            "Försöker förstå sig själv och sina känslor, efterfrågar feedback i detta hänseende och är uppmärksam på sin inverkan på andra. "
        ),
        "Uthållighet": (
            "Fungerar bra under press och kommer snabbt igen efter motgångar på ett positivt sätt. "
        ),
    },
    "strategiskt_tankande_och_anpassningsformaga": {
        "Strategiskt fokus": (
            "Visar prov på ett strategiskt förhållningssätt i sitt arbete, tar hänsyn till hur olika aspekter av organisationen interagerar, helhetsbilden och företagets framtid på längre sikt. "
        ),
        "Anpassningsförmåga": (
            "Anpassar sitt tillvägagångssätt och reagerar effektivt på olika situationer, människor och möjligheter. "
        ),
    },
    "kommunikation_och_samarbete": {
        "Teamarbete": (
            "Arbetar tillsammans med teamet för att uppnå gemensamma mål, lösa potentiella utmaningar och främja ett gemensamt agerande."
        ),
        "Inflytelserik": (
            "Påverkar andras handlingar och åsikter med hjälp av övertygande argument och strategier. "
        ),
    },
}


SECTION_AI_CONFIG = [
    {
        "prompt_name": "leda",
        "section_key": "leda_utveckla_och_engagera",
        "subscales": [
            "Leda andra",
            "Engagera andra",
            "Delegera",
            "Utveckla andra",
        ],
        "label": "Leda, utveckla och engagera",
        "context_key": "leda_text",
        "table_context_key": "leda_table_html",
    },
    {
        "prompt_name": "mod",
        "section_key": "mod_och_handlingskraft",
        "subscales": [
            "Beslutsamhet",
            "Integritet",
            "Hantera konflikter",
        ],
        "label": "Mod och handlingskraft",
        "context_key": "mod_text",
        "table_context_key": "mod_table_html",
    },
    {
        "prompt_name": "sjalkannedom",
        "section_key": "sjalkannedom_och_emotionell_stabilitet",
        "subscales": [
            "Självmedvetenhet",
            "Uthållighet",
        ],
        "label": "Självkännedom och emotionell stabilitet",
        "context_key": "sjalkannedom_text",
        "table_context_key": "sjalkannedom_table_html",
    },
    {
        "prompt_name": "strategi",
        "section_key": "strategiskt_tankande_och_anpassningsformaga",
        "subscales": [
            "Strategiskt fokus",
            "Anpassningsförmåga",
        ],
        "label": "Strategiskt tänkande och anpassningsförmåga",
        "context_key": "strategi_text",
        "table_context_key": "strategi_table_html",
    },
    {
        "prompt_name": "kommunikation",
        "section_key": "kommunikation_och_samarbete",
        "subscales": [
            "Teamarbete",
            "Inflytelserik",
        ],
        "label": "Kommunikation och samarbete",
        "context_key": "kommunikation_text",
        "table_context_key": "kommunikation_table_html",
    },
]

# För snabb lookup: "leda andra" -> ("leda_utveckla_och_engagera", "Leda andra")
LABEL_TO_TARGET = {}
for section, subs in TARGETS.items():
    for sub in subs:
        LABEL_TO_TARGET[sub.lower()] = (section, sub)


def _map_0_10_to_1_5(x) -> int:
    """Mappa 0–10 (eller 1–10) till 1–5."""
    try:
        v = float(str(x).replace(",", "."))
    except Exception:
        return 3
    # tillåt 0–10 och clamp:a
    if v < 0:
        v = 0
    if v > 10:
        v = 10
    # intervall om 2 poäng: 0-1.99=>1, 2-3.99=>2, 4-5.99=>3, 6-7.99=>4, 8-10=>5
    bucket = 1 + int(v // 2.0)
    if bucket > 5:
        bucket = 5
    if bucket < 1:
        bucket = 1
    return bucket

def html_to_text(html: str) -> str:
    """
    Tar HTML och returnerar ren text i ett Word-vänligt format.
    Behåller radbrytningar, listor m.m.
    """
    if not html:
        return ""

    soup = BeautifulSoup(html, "html.parser")

    # <br> -> radbrytning
    for br in soup.find_all("br"):
        br.replace_with("\n")

    lines = []

    for elem in soup.recursiveChildGenerator():
        if elem.name == "li":
            # punktlista
            text = elem.get_text(separator=" ", strip=True)
            if text:
                lines.append("• " + text)

        elif elem.name in ("p", "div"):
            # viktig ändring: separator="\n" gör att radbrytningar inne i <p> bevaras
            text = elem.get_text(separator="\n", strip=True)
            if text:
                lines.append(text)

    # Om vi inte hittade några p/div/li: fallback
    if not lines:
        return soup.get_text("\n", strip=True)

    # En tom rad mellan block kan du få med "\n\n" om du vill ha mer luft.
    return "\n".join(lines)


def _normalize_header_cell(value: str) -> str:
    """
    Tar t.ex. 'Competency Score: Directing others (STIVE)'
    -> 'directing others'
    (lowercase + trimmad, utan 'Competency Score:' och '(STIVE)')
    """
    if not value:
        return ""
    text = str(value)
    if "Competency Score" in text:
        text = text.split("Competency Score:")[-1]
    if "(" in text:
        text = text.split("(")[0]
    return text.strip().lower()


def _ratings_from_worksheet(ws):
    """
    STIVE-format:
      - Rad 1: rubriker ('Competency Score: Leading others (STIVE)')
      - Rad 2: värden (1–5, ev. decimal)
    Mappar med HEADER_TO_TARGET till svenska etiketter per sektion.
    """
    rows = list(ws.iter_rows(values_only=True))
    debug = []

    if len(rows) < 2:
        debug.append("Excel behöver minst två rader (rubriker + en rad med resultat).")
        return _default_all_three(), debug

    header = rows[0]
    data = rows[1]

    ratings = _default_all_three()

    for col_idx, raw_header in enumerate(header):
        # hoppa över förnamn/efternamn om de ligger först
        if col_idx < 2:
            continue

        comp_name = _normalize_header_cell(raw_header)
        if not comp_name:
            continue

        mapping = HEADER_TO_TARGET.get(comp_name)
        if not mapping:
            debug.append(f"Kolumn {col_idx+1}: '{raw_header}' ignoreras (ingen mapping).")
            continue

        sec_key, sub_label = mapping

        raw_value = data[col_idx] if col_idx < len(data) else None
        if raw_value is None or raw_value == "":
            debug.append(f"{comp_name}: inget värde i rad 2, behåller default.")
            continue

        score = _round_to_1_5(raw_value)
        score = max(1, min(5, score))

        if sec_key not in ratings:
            ratings[sec_key] = {}
        ratings[sec_key][sub_label] = score

        debug.append(f"{comp_name} -> {sub_label}: rå={raw_value} -> {score}")

    return ratings, debug

DATA_URL_RE = re.compile(r"^data:image\/[a-zA-Z]+;base64,")

IMAGE_WIDTHS_IN = {
    "{leda_image}": 4.6,
    "{mod_image}": 4.6,
    "{sjalkannedom_image}": 4.6,
    "{strategi_image}": 4.6,
    "{kommunikation_image}": 4.6,
}



def replace_image_placeholder(doc, placeholder: str, data_url: str, width_in: float | None = None):
    """
    Ersätter placeholder i Word med en bild (base64 dataURL).
    width_in: om None -> hämtas från IMAGE_WIDTHS_IN, fallback 5.8
    """
    if not data_url:
        return

    b64 = DATA_URL_RE.sub("", data_url).strip()
    if not b64:
        return

    try:
        img_bytes = base64.b64decode(b64)
    except Exception:
        return

    target_width = width_in if width_in is not None else IMAGE_WIDTHS_IN.get(placeholder, 5.8)

    def _insert_into_paragraph(p):
        if placeholder not in (p.text or ""):
            return False

        # töm runs
        for r in p.runs:
            r.text = ""

        img_stream = BytesIO(img_bytes)
        img_stream.seek(0)

        run = p.add_run()
        run.add_picture(img_stream, width=Inches(target_width))
        return True

    # 1) vanliga paragrafer
    for p in doc.paragraphs:
        if _insert_into_paragraph(p):
            return

    # 2) tabeller
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _insert_into_paragraph(p):
                        return


def docx_replace_text(doc, mapping: dict):
    """
    Ersätter t.ex. {candidate_name} i alla paragrafer och tabellceller.
    Enkel variant som funkar bra om du inte har massa blandad formatering
    mitt i samma mening.
    """
    def replace_in_paragraph(paragraph):
        if not paragraph.text:
            return
        full_text = paragraph.text
        changed = False
        for k, v in mapping.items():
            if k in full_text:
                full_text = full_text.replace(k, v or "")
                changed = True
        if changed:
            # nollställ runs och skriv allt i första run
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = full_text

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

TABLE_PLACEHOLDERS = {
    "{leda_table}": "leda_utveckla_och_engagera",
    "{mod_table}": "mod_och_handlingskraft",
    "{sjalkannedom_table}": "sjalkannedom_och_emotionell_stabilitet",
    "{strategi_table}": "strategiskt_tankande_och_anpassningsformaga",
    "{kommunikation_table}": "kommunikation_och_samarbete",
}

def _insert_ratings_table_into_cell(cell, section_key, ratings_dict):
    """
    Skapar en liten 1–5-tabell inne i en cell, baserat på ratings_json.
    - cell: docx.cell
    - section_key: t.ex. 'leda_utveckla_och_engagera'
    - ratings_dict: dict från ratings_json
    """
    section_ratings = ratings_dict.get(section_key, {})
    labels = TARGETS.get(section_key, [])
    if not labels:
        return

    # skapa inre tabell: 1 kolumn för label + 5 för prickar
    inner = cell.add_table(rows=len(labels), cols=6)
    inner.style = "Table Grid"  # eller någon stil du har i mallen

    for r_idx, label in enumerate(labels):
        row = inner.rows[r_idx]
        # första kolumnen = etikett
        row.cells[0].text = label

        score = int(section_ratings.get(label, 3)) if section_ratings.get(label) is not None else 3
        score = max(1, min(5, score))

        # kolumn 1–5 = ○/●
        for c_idx in range(1, 6):
            cell_dot = row.cells[c_idx]
            p = cell_dot.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("●" if c_idx == score else "○")
            run.font.size = Pt(10)


def _apply_table_placeholders(doc, ratings_dict):
    """
    Går igenom alla tabell-taggar ({leda_table} osv) och ersätter dem
    med inre tabeller baserade på ratings_dict.
    """
    if not ratings_dict:
        return

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text or ""
                    for placeholder, section_key in TABLE_PLACEHOLDERS.items():
                        if placeholder in text:
                            # ta bort taggen i paragrafen
                            p.text = text.replace(placeholder, "").strip()
                            # stoppa in vår inre tabell i cellen
                            _insert_ratings_table_into_cell(cell, section_key, ratings_dict)

def _safe_table_style(table, style_name="Table Grid"):
    """
    Försök sätta en tabellstil om den finns i dokumentmallen.
    Annars: låt Word använda standardstilen.
    """
    try:
        table.style = style_name
    except KeyError:
        # Stilen finns inte i denna .dotx/.docx → använd default
        pass

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def build_ratings_table_for_section(doc, ratings: dict, section_key: str):
    """
    Skapar en “snygg” tabell för en given sektion, t.ex. leda_utveckla_och_engagera.

    | Leda andra
      (beskrivning)             ○  ●  ○  ○  ○ |

    Första kolumnen: etikett + kort beskrivning.
    Nästa 5 kolumner: cirklar där rätt värde (1–5) är ifyllt.
    """
    section_ratings = ratings.get(section_key) or {}
    labels_order = TARGETS.get(section_key) or list(section_ratings.keys())

    rows = len(labels_order)
    cols = 6  # label + 5 cirklar

    # Skapa tabell – den flyttas med replace_table_placeholder till rätt ställe
    table = doc.add_table(rows=rows, cols=cols)
    _safe_table_style(table, "Table Grid")  # försöker sätta stil om den finns

    for r_idx, label in enumerate(labels_order):
        row = table.rows[r_idx]

        # ---------- Första cellen: rubrik + beskrivning ----------
        label_cell = row.cells[0]

        # rubrik
        p_label = label_cell.paragraphs[0]
        run_label = p_label.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(10)

        # beskrivning (om vi har en)
        desc = SUBSCALE_DESCRIPTIONS.get(section_key, {}).get(label)
        if desc:
            p_desc = label_cell.add_paragraph(desc)
            p_desc.paragraph_format.space_before = 0
            p_desc.paragraph_format.space_after = 0
            for run in p_desc.runs:
                run.font.size = Pt(8)

        # ---------- Poäng: 1–5 cirklar ----------
        raw_score = section_ratings.get(label, 3)
        try:
            score = int(raw_score)
        except Exception:
            score = 3
        score = max(1, min(5, score))

        for c in range(1, 6):
            cell = row.cells[c]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("●" if c == score else "○")
            run.font.size = Pt(9)

    return table


def replace_table_placeholder(doc, placeholder: str, ratings: dict, section_key: str):
    """
    Hittar paragrafen som innehåller t.ex. {leda_table},
    lägger in tabellen direkt efter paragrafen (i samma cell eller i body),
    och tar bort själva taggen.
    """

    def _handle_paragraph(p):
        if placeholder not in p.text:
            return False

        # Skapa tabellen (först i dokumentets body)
        table = build_ratings_table_for_section(doc, ratings, section_key)

        # Flytta tabellen så att den hamnar direkt efter paragrafen p
        p._p.addnext(table._tbl)

        # Ta bort taggen ur paragrafens text
        for run in p.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, "")
        return True

    # 1) Sök i body-paragrafer
    for p in doc.paragraphs:
        if _handle_paragraph(p):
            return

    # 2) Sök i alla tabellceller
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _handle_paragraph(p):
                        return

# ── NYTT: statisk skalförklaring (HTML) med header ───────────────────────────
def _scale_demo_html() -> str:
    demo = {
        "leda_utveckla_och_engagera": {"Exempel": 3},
        "mod_och_handlingskraft": {"Exempel": 3},
        "sjalkannedom_och_emotionell_stabilitet": {"Exempel": 3},
        "strategiskt_tankande_och_anpassningsformaga": {"Exempel": 3},
        "kommunikation_och_samarbete": {"Exempel": 3},
    }
    return _ratings_table_html(demo)


# ──────────────────────────────────────────────────────────────────────────────
# Defaults: skapas per användare om inget finns
# ──────────────────────────────────────────────────────────────────────────────
def ensure_default_prompts_exist(user):
    defaults = {
        "global_style": (
            "❌ Använd aldrig taggar i global_style.\n\n"
            "Global stil ska ENDAST innehålla:\n"
            "- skrivregler\n"
            "- tonalitet\n"
            "- språkbruk (t.ex. “säg grupp, inte team”)\n"
            "- strukturella regler (t.ex. “undvik punktlistor”, “skriv i tredje person”)\n\n"
            "Den får inte innehålla innehållstexter eller taggar som "
            "{excel_text}, {intervju_text}, {uploaded_files}."
        ),
        # NYTT — tolkning av Excelfil
        "tolka_excel_resultat": (
            "Du är en analytiker som ska tolka en Excelfil med testresultat. "
            "Varje kolumn representerar en kompetens och en poäng mellan 1 och 5. "
            "Din uppgift är att identifiera varje kompetens och tilldela ett betyg mellan 1 och 5. "
            "Om du ser decimaltal, avrunda 2,2 till 2 och 2,7 till 3. "
            "Returnera endast JSON i följande format:\n\n"
            "{\n"
            '  "Leda, utveckla och engagera": {\n'
            '    "Leda andra": 3,\n'
            '    "Engagera andra": 4,\n'
            '    "Delegera": 3,\n'
            '    "Utveckla andra": 2\n'
            "  },\n"
            '  "Mod och handlingskraft": {\n'
            '    "Beslutsamhet": 4,\n'
            '    "Integritet": 5,\n'
            '    "Hantera konflikter": 3\n'
            "  }\n"
            "}\n\n"
            "Skriv *endast* JSON, ingen förklarande text."
        ),

        "betygsskala_forklaring": (
            "1 = Tydligt utvecklingsområde. Beteendet stödjer inte kraven i rollen.\n"
            "2 = Acceptabel nivå, men med tydliga utvecklingsbehov i mer komplexa situationer.\n"
            "3 = God nivå. Räcker för de flesta vardagliga krav.\n"
            "4 = Stark nivå. Personen visar ofta beteenden som stödjer rollen väl.\n"
            "5 = Mycket stark nivå. Personen ligger konsekvent högt och fungerar som förebild.\n\n"
            "Specifika beskrivningar per kompetens:\n"
            "- Leda andra: 1 = Den här personen är..., 3 = Den här personen är..., 5 = Den här personen är...\n"
            "- Beslutsamhet: 1 = …, 3 = …, 5 = …\n"
        ),

        # befintliga
        #"testanalys": """Du är en psykolog specialiserad på testtolkning...
#{excel_text}
#""",
        #"intervjuanalys": """Du är en HR-expert...
#{intervjuanteckningar}
#""",
        #"helhetsbedomning": """Du är en HR-expert...
#Test:
#{test_text}

#Intervju:
#{intervju_text}
#""",

        # per-rubrik
        "tq_fardighet": "Skriv TQ Färdighet baserat på testdata.\n\n{ratings_json}\n\n{intervju_text}",
        "tq_motivation": "Identifiera de tre främsta motivationsfaktorerna och beskriv kort.\n\n{ratings_json}\n\n{intervju_text}\n\n{betygsskala_forklaring}",
        "leda": "Skriv bedömning för 'Leda, utveckla och engagera' med fokus på testdata och komplettera med intervju.\n\n{excel_text}\n\n{intervju_text}\n\n{betygsskala_forklaring}",
        "mod": "Skriv bedömning för 'Mod och handlingskraft'.\n\n{ratings_json}\n\n{intervju_text}\n\n{betygsskala_forklaring}",
        "sjalkannedom": "Skriv bedömning för 'Självkännedom och emotionell stabilitet'.\n\n{ratings_json}\n\n{intervju_text}\n\n{betygsskala_forklaring}",
        "strategi": "Skriv bedömning för 'Strategiskt tänkande och anpassningsförmåga'.\n\n{ratings_json}\n\n{intervju_text}\n\n{betygsskala_forklaring}",
        "kommunikation": "Skriv bedömning för 'Kommunikation och samarbete'.\n\n{ratings_json}\n\n{intervju_text}\n\n{betygsskala_forklaring}",

        # sammanställningar
        "styrkor_utveckling_risk": (
            "Sammanfatta till tre listor: Styrkor, Utvecklingsområden, Riskbeteenden. "
            "Använd de sju sektionerna nedan som källa.\n\n"
            "TQ Färdighet:\n{tq_fardighet_text}\n\n"
            "TQ Motivation:\n{tq_motivation_text}\n\n"
            "Leda:\n{leda_text}\n\nMod:\n{mod_text}\n\nSjälvkännedom:\n{sjalkannedom_text}\n\n"
            "Strategi:\n{strategi_text}\n\nKommunikation:\n{kommunikation_text}"
        ),
        "sammanfattande_slutsats": (
            "Skriv en sammanfattande slutsats (1–2 stycken) som väger samman allt. "
            "Ta hänsyn till Styrkor/Utvecklingsområden/Risk och alla sektioner.\n\n"
            "Styrkor/Utvecklingsområden/Risk:\n{sur_text}\n\n"
            "TQ Färdighet:\n{tq_fardighet_text}\n\n"
            "TQ Motivation:\n{tq_motivation_text}\n\n"
            "Leda:\n{leda_text}\n\nMod:\n{mod_text}\n\nSjälvkännedom:\n{sjalkannedom_text}\n\n"
            "Strategi:\n{strategi_text}\n\nKommunikation:\n{kommunikation_text}"
        ),
    }

    for name, text in defaults.items():
        Prompt.objects.get_or_create(user=user, name=name, defaults={"text": text})



# ──────────────────────────────────────────────────────────────────────────────
# Hjälpare
# ──────────────────────────────────────────────────────────────────────────────
def _trim(s: str, max_chars: int = 6500) -> str:
    """Trimma långa texter (behåll början och slut) för att undvika tokenproblem."""
    s = s or ""
    if len(s) <= max_chars:
        return s
    head = s[: max_chars // 2]
    tail = s[- max_chars // 2 :]
    return head + "\n...\n" + tail

def _safe_json_from_text(txt: str):
    """
    Försök plocka JSON efter rubriken RATINGS_JSON.
    Hanterar även ```json ... ``` och whitespace.
    """
    if not txt:
        return None
    m = re.search(r"###\s*RATINGS_JSON\s*(.*)$", txt, flags=re.DOTALL | re.IGNORECASE)
    if not m:
        return None
    block = m.group(1).strip()

    fence = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", block, flags=re.DOTALL)
    if fence:
        block = fence.group(1)
    else:
        brace = re.search(r"(\{.*\})", block, flags=re.DOTALL)
        if brace:
            block = brace.group(1)

    try:
        return json.loads(block)
    except Exception:
        return None

def _default_all_three():
    """Sista fallback — fyll 3:or överallt så UI alltid renderar."""
    return {
        "leda_utveckla_och_engagera": {
            "Leda andra": 3, "Engagera andra": 3, "Delegera": 3, "Utveckla andra": 3
        },
        "mod_och_handlingskraft": {
            "Beslutsamhet": 3, "Integritet": 3, "Hantera konflikter": 3
        },
        "sjalkannedom_och_emotionell_stabilitet": {
            "Självmedvetenhet": 3, "Uthållighet": 3
        },
        "strategiskt_tankande_och_anpassningsformaga": {
            "Strategiskt fokus": 3, "Anpassningsförmåga": 3
        },
        "kommunikation_och_samarbete": {
            "Teamarbete": 3, "Inflytelserik": 3
        }
    }

def _markdown_to_html(text: str) -> str:
    if not text:
        return ""

    text = str(text).strip()

    # 🔍 Om texten redan innehåller typiska HTML-taggar: använd som den är
    # (detta händer när användaren har redigerat i WYSIWYG och vi får tillbaka HTML)
    if "<" in text and re.search(r"</(p|br|strong|em|ul|ol|li|h[1-6])\s*>", text, flags=re.IGNORECASE):
        return mark_safe(text)

    # Annars: behandla det som markdown från OpenAI
    html = markdown(
        text,
        extras=[
            "fenced-code-blocks",
            "tables",
            "strike",
            "break-on-newline",   # 👈 gör \n till radbrytningar
            "smarty-pants",
            "spoiler",
            "header-ids",
            "cuddled-lists",
        ]
    )
    return mark_safe(html)

# ──────────────────────────────────────────────────────────────────────────────
# Prompt Editor (om du har en sida för att redigera prompter)
# ──────────────────────────────────────────────────────────────────────────────
@login_required
@csrf_exempt
def prompt_editor(request):
    owner = get_prompt_owner(request.user)
    ensure_default_prompts_exist(owner)

    sidebar_session, _ = ChatSession.objects.get_or_create(
        user=request.user,
        title="Domarnämnden-verktygschatt",
    )

    prompts = Prompt.objects.filter(user=owner)

    # ✅ Bara prompt-ägaren (t.ex. Veronika) får ändra
    can_edit = (request.user == owner)

    if request.method == "POST" and can_edit:
        # reset-logik kan vara kvar om du vill
        if "reset" in request.POST:
            name = request.POST["reset"]
            defaults = {
                "testanalys": """Du är en psykolog specialiserad på testtolkning...""",
                "intervjuanalys": """Du är en HR-expert. Nedan finns intervjuanteckningar...""",
                "helhetsbedomning": """Du är en HR-expert. Nedan finns en testanalys..."""
            }
            if name in defaults:
                prompt = Prompt.objects.get(user=owner, name=name)
                prompt.text = defaults[name]
                prompt.save()
        else:
            for prompt in prompts:
                field_name = f"prompt_{prompt.name}"
                new_text = request.POST.get(field_name)
                if new_text is not None:
                    prompt.text = new_text
                    prompt.save()

    return render(
        request,
        "prompt_editor.html",
        {
            "prompts": prompts,
            "can_edit": can_edit,   # 👈 skickas till templaten
            "owner": owner,
        },
    )

# ──────────────────────────────────────────────────────────────────────────────
# Assistent-sidofält
# ──────────────────────────────────────────────────────────────────────────────


SECTION_TITLES = {
    "leda_utveckla_och_engagera": "Leda, utveckla och engagera",
    "mod_och_handlingskraft": "Mod och handlingskraft",
    "sjalkannedom_och_emotionell_stabilitet": "Självkännedom och emotionell stabilitet",
    "strategiskt_tankande_och_anpassningsformaga": "Strategiskt tänkande och anpassningsförmåga",
    "kommunikation_och_samarbete": "Kommunikation och samarbete",
}

def _build_sidebar_ratings(ratings: dict):
    """
    Gör om ratings-dict till något som är enkelt att loopa i templaten:
    {
      "Leda, utveckla och engagera": [
        {"label": "Leda andra", "value": 5},
        {"label": "Engagera andra", "value": 4},
        ...
      ],
      ...
    }
    """
    grouped = {}

    for sec_key, title in SECTION_TITLES.items():
        sec = ratings.get(sec_key)
        if not sec:
            continue

        grouped[title] = []

        # Försök hålla ordning enligt TARGETS om det finns
        labels_order = TARGETS.get(sec_key) or list(sec.keys())

        for label in labels_order:
            val = sec.get(label)
            if val is None:
                continue
            try:
                v = int(val)
            except Exception:
                v = val
            grouped[title].append({
                "label": label,
                "value": v,
            })

    return grouped

# ──────────────────────────────────────────────────────────────────────────────
# Huvudvy
# ──────────────────────────────────────────────────────────────────────────────
@login_required
@csrf_exempt
def index(request):
    owner = get_prompt_owner(request.user)
    ensure_default_prompts_exist(owner)

    # ---------- 1) Läs nuvarande steg (GET eller POST) ----------
    try:
        step = int(request.POST.get("step", request.GET.get("step", "1")))
    except ValueError:
        step = 1

    # --- NYTT (FIX): Ladda report_id / report-data tidigt så loaded_data kan användas i steglogiken ---
    report_id = request.GET.get("report_id") or request.POST.get("report_id")
    loaded_report = None
    loaded_data = {}

    report_id = (request.GET.get("report_id") or request.POST.get("report_id") or "").strip()
    if report_id.lower() in ("none", "null", ""):
        report_id = ""

    if report_id:
        loaded_report = _get_report_or_404(report_id)
        loaded_data = loaded_report.data or {}
    
    edit_mode = bool(report_id and loaded_data)

    # 🔹 RESET när man kommer in "från början" (GET på steg 1)
    if request.method == "GET" and step == 1 and not edit_mode:
        request.session.pop("selected_motivation_keys", None)  # FIX: bara en gång
        request.session.pop("motivation_notes", None)
        request.session.pop("job_ad_text", None)
        request.session.pop("logical_score", None)
        request.session.pop("verbal_score", None)
        # om du i framtiden sparar fler saker i sessionen för rapporten
        # kan du tömma dem här också med fler .pop(...)
        # request.session.pop("some_other_key", None)

    # 🔹 Ladda ev. sparade motivationsval från session
    selected_motivation_keys = request.session.get("selected_motivation_keys", [])

    # 🔹 NYTT: plocka ut namnvarianter och normalisera
    raw_first = (request.POST.get("candidate_first_name") or "").strip()
    raw_last  = (request.POST.get("candidate_last_name") or "").strip()
    raw_full  = (request.POST.get("candidate_name") or "").strip()

    full_name = (raw_first + " " + raw_last).strip()

    # ---------- 2) Plocka in state från POST ----------
    context = {
        "step": step,
        "test_text": request.POST.get("test_text", ""),
        "intervju_text": request.POST.get("intervju_text", ""),
        "tq_fardighet_text": request.POST.get("tq_fardighet_text", ""),
        "tq_motivation_text": request.POST.get("tq_motivation_text", ""),
        "leda_text": request.POST.get("leda_text", ""),
        "mod_text": request.POST.get("mod_text", ""),
        "sjalkannedom_text": request.POST.get("sjalkannedom_text", ""),
        "strategi_text": request.POST.get("strategi_text", ""),
        "kommunikation_text": request.POST.get("kommunikation_text", ""),
        "sur_text": request.POST.get("sur_text", ""),
        "slutsats_text": request.POST.get("slutsats_text", ""),
        "cv_text": request.POST.get("cv_text", ""),
        "selected_motivation_keys": selected_motivation_keys,

        # 🔹 extra inputs
        "job_ad_text": request.POST.get("job_ad_text") or request.session.get("job_ad_text", ""),
        "motivation_notes": request.POST.get("motivation_notes") or request.session.get("motivation_notes", ""),
        "logical_score": request.POST.get("logical_score") or request.session.get("logical_score", ""),
        "verbal_score": request.POST.get("verbal_score") or request.session.get("verbal_score", ""),

        # kandidatinfo
        "candidate_first_name": raw_first,
        "candidate_last_name": raw_last,
        "candidate_name": request.POST.get("candidate_name", ""),
        "candidate_role": request.POST.get("candidate_role", ""),

        # CV som markdown/HTML
        "uploaded_files_markdown": request.POST.get("uploaded_files_markdown", ""),
        "uploaded_files_html": "",

        "error": "",
    }

    # Om markdown finns i POST → skapa HTML igen
    if context["uploaded_files_markdown"]:
        context["uploaded_files_html"] = markdown(context["uploaded_files_markdown"])

    # ✅ När report_id finns: fyll context från sparad report.data (GET + POST)
    if report_id and loaded_data:
        context = _apply_report_data_to_context(context, loaded_data)

        # synka sessionen (valfritt men bra för sidopanelen)
        request.session["selected_motivation_keys"] = context.get("selected_motivation_keys", []) or []
        request.session["job_ad_text"] = context.get("job_ad_text", "") or ""
        request.session["motivation_notes"] = context.get("motivation_notes", "") or ""
        request.session["logical_score"] = context.get("logical_score", "") or ""
        request.session["verbal_score"] = context.get("verbal_score", "") or ""
        request.session.modified = True

    # ---------- 3) Ratings (JSON + tabeller) ----------
    ratings_json_str = (
        request.POST.get("ratings_json")
        or context.get("ratings_json")   # från loaded_data via apply_report_data
        or ""
    )
    ratings = None
    if ratings_json_str:
        try:
            ratings = json.loads(ratings_json_str)
        except Exception:
            ratings = None

    if ratings:
        context["ratings_json"] = ratings_json_str

        context["leda_table_html"] = mark_safe(_ratings_table_html(
            ratings,
            section_filter=[("leda_utveckla_och_engagera", "Leda, utveckla och engagera")],
            include_css=True,
        ))
        context["mod_table_html"] = mark_safe(_ratings_table_html(
            ratings,
            section_filter=[("mod_och_handlingskraft", "Mod och handlingskraft")],
            include_css=True,
        ))
        context["sjalkannedom_table_html"] = mark_safe(_ratings_table_html(
            ratings,
            section_filter=[("sjalkannedom_och_emotionell_stabilitet", "Självkännedom och emotionell stabilitet")],
            include_css=True,
        ))
        context["strategi_table_html"] = mark_safe(_ratings_table_html(
            ratings,
            section_filter=[("strategiskt_tankande_och_anpassningsformaga", "Strategiskt tänkande och anpassningsförmåga")],
            include_css=True,
        ))
        context["kommunikation_table_html"] = mark_safe(_ratings_table_html(
            ratings,
            section_filter=[("kommunikation_och_samarbete", "Kommunikation och samarbete")],
            include_css=True,
        ))
        context["ratings_sidebar"] = _build_sidebar_ratings(ratings)

    # ---------- 3.5) Motivationsfaktorer till sidopanelen ----------
    selected_keys = request.session.get("selected_motivation_keys", [])

    selected_motivations = []
    for key in selected_keys:
        data = MOTIVATION_FACTORS.get(key)
        if not data:
            continue
        selected_motivations.append({
            "key": key,
            "label": data["label"],
            "definition": data["definition"],
        })

    # FIX: sätt i context efter loopen (istället för inne i loopen)
    context["selected_motivations"] = selected_motivations

    # --- NYTT (FIX): bygg lista med valda motivationsfaktorer att skicka till AI ---
    selected_motivations_for_ai = [
        {
            "key": k,
            "label": MOTIVATION_FACTORS[k]["label"],
            "definition": MOTIVATION_FACTORS[k]["definition"],
        }
        for k in selected_keys
        if k in MOTIVATION_FACTORS
    ]

    # ---------- 4) POST-actions (prev/next/build_doc) ----------
    if request.method == "POST":

        # Föregående
        if "prev" in request.POST:
            step = max(1, step - 1)

        # Skapa Word endast på sista steget
        elif "build_doc" in request.POST and step == 11:
            from django.http import HttpResponse
            from docx import Document

            template_path = os.path.join(
                settings.BASE_DIR,
                "reports",
                "domarnamnden_template.docx"
            )
            doc = Document(template_path)

            leda_image_data = request.POST.get("leda_image", "")
            mod_image_data = request.POST.get("mod_image", "")
            sjalkannedom_image_data = request.POST.get("sjalkannedom_image", "")
            strategi_image_data = request.POST.get("strategi_image", "")
            kommunikation_image_data = request.POST.get("kommunikation_image", "")

            # --- bygg text för valda motivationsfaktorer ---
            selected_motivations_doc = context.get("selected_motivations") or []

            motivation_lines = []
            for mot in selected_motivations_doc:
                if isinstance(mot, dict):
                    label = mot.get("label", "")
                    definition = mot.get("definition", "")
                else:
                    label = getattr(mot, "label", "")
                    definition = getattr(mot, "definition", "")

                if label or definition:
                    motivation_lines.append(f"{label}\n{definition}".strip())

            selected_motivations_text = "\n\n".join(motivation_lines)

            mapping = {
                "{candidate_name}": context.get("candidate_name", ""),
                "{candidate_first_name}": context.get("candidate_first_name", ""),
                "{candidate_last_name}": context.get("candidate_last_name", ""),
                "{candidate_role}": context.get("candidate_role", ""),
                "{tq_fardighet_text}": html_to_text(context.get("tq_fardighet_text", "")),
                "{sur_text}": html_to_text(context.get("sur_text", "")),
                "{tq_motivation_text}": html_to_text(context.get("tq_motivation_text", "")),
                "{leda_text}": html_to_text(context.get("leda_text", "")),
                "{mod_text}": html_to_text(context.get("mod_text", "")),
                "{sjalkannedom_text}": html_to_text(context.get("sjalkannedom_text", "")),
                "{strategi_text}": html_to_text(context.get("strategi_text", "")),
                "{kommunikation_text}": html_to_text(context.get("kommunikation_text", "")),
                "{selected_motivations}": selected_motivations_text,
            }
            docx_replace_text(doc, mapping)

            ratings_json_raw = (
                request.POST.get("ratings_json")
                or context.get("ratings_json")
                or ""
            )

            ratings_doc = {}
            if isinstance(ratings_json_raw, dict):
                ratings_doc = ratings_json_raw
            elif isinstance(ratings_json_raw, str) and ratings_json_raw.strip():
                try:
                    ratings_doc = json.loads(ratings_json_raw)
                except json.JSONDecodeError:
                    ratings_doc = {}

            if ratings_doc:
                replace_table_placeholder(
                    doc, "{leda_table}", ratings_doc, "leda_utveckla_och_engagera"
                )
                replace_table_placeholder(
                    doc, "{mod_table}", ratings_doc, "mod_och_handlingskraft"
                )
                replace_table_placeholder(
                    doc, "{sjalkannedom_table}", ratings_doc,
                    "sjalkannedom_och_emotionell_stabilitet"
                )
                replace_table_placeholder(
                    doc, "{strategi_table}", ratings_doc,
                    "strategiskt_tankande_och_anpassningsformaga"
                )
                replace_table_placeholder(
                    doc, "{kommunikation_table}", ratings_doc,
                    "kommunikation_och_samarbete"
                )

            replace_image_placeholder(doc, "{leda_image}", leda_image_data)
            replace_image_placeholder(doc, "{mod_image}", mod_image_data)
            replace_image_placeholder(doc, "{sjalkannedom_image}", sjalkannedom_image_data)
            replace_image_placeholder(doc, "{strategi_image}", strategi_image_data)
            replace_image_placeholder(doc, "{kommunikation_image}", kommunikation_image_data)

            response = HttpResponse(
                content_type=(
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document"
                )
            )
            filename = f"bedomning_{context.get('candidate_name','rapport')}.docx"
            response["Content-Disposition"] = f'attachment; filename=\"{filename}\"'
            doc.save(response)
            return response

        # Nästa (inkl AI)
        elif "next" in request.POST:

            try:
                style = Prompt.objects.get(user=owner, name="global_style").text
            except Prompt.DoesNotExist:
                style = getattr(settings, "STYLE_INSTRUCTION", "")

            try:
                betygsskala_prompt = Prompt.objects.get(
                    user=owner, name="betygsskala_forklaring"
                ).text
            except Prompt.DoesNotExist:
                betygsskala_prompt = ""

            ratings_json_str = (
                request.POST.get("ratings_json")
                or context.get("ratings_json")   # från loaded_data via apply_report_data
                or ""
            )

            # ---------- STEG 1 ----------
            if step == 1:
                excel_text = ""
                ws = None

                # Läs för- och efternamn från formuläret
                first = (request.POST.get("candidate_first_name") or "").strip()
                last  = (request.POST.get("candidate_last_name") or "").strip()
                role  = (request.POST.get("candidate_role") or "").strip()

                full_name = (first + " " + last).strip()

                # uppdatera context så det följer med vidare
                context["candidate_first_name"] = first
                context["candidate_last_name"] = last
                context["candidate_name"] = full_name
                context["candidate_role"] = role

                if not first or not last:
                    context["error"] = "Fyll i både förnamn och efternamn."

                # jobbannons från PDF
                job_ad_text = ""
                job_ad_file = request.FILES.get("job_ad_pdf")

                if job_ad_file:
                    job_ad_text = _trim(_read_file_text(job_ad_file), 6500).strip()
                    if not job_ad_text:
                        context["error"] = "Kunde inte läsa någon text från jobbannons-PDF:en."
                else:
                    job_ad_text = ""

                context["job_ad_text"] = job_ad_text
                request.session["job_ad_text"] = job_ad_text

                # jobbannons + motivationsanteckningar
                motivation_notes = (request.POST.get("motivation_notes") or "").strip()
                context["job_ad_text"] = job_ad_text
                context["motivation_notes"] = motivation_notes

                request.session["job_ad_text"] = job_ad_text
                request.session["motivation_notes"] = motivation_notes

                # Hämta valda motivationsfaktorer (max 3)
                motivation_choices = request.POST.getlist("motivation_choices")
                motivation_choices = motivation_choices[:3]

                context["selected_motivation_keys"] = motivation_choices
                request.session["selected_motivation_keys"] = motivation_choices

                # FIX: uppdatera även AI-listan efter att valen kommit in
                selected_motivations_for_ai = [
                    {
                        "key": k,
                        "label": MOTIVATION_FACTORS[k]["label"],
                        "definition": MOTIVATION_FACTORS[k]["definition"],
                    }
                    for k in motivation_choices
                    if k in MOTIVATION_FACTORS
                ]

                # färdighetsvärden 0–99
                logical_raw = (request.POST.get("logical_score") or "").strip()
                verbal_raw = (request.POST.get("verbal_score") or "").strip()
                context["logical_score"] = logical_raw
                context["verbal_score"] = verbal_raw

                def _validate_0_99(label, raw):
                    if not raw:
                        return None, f"Ange en siffra 0–99 för {label}."
                    try:
                        val = int(raw)
                    except ValueError:
                        return None, f"Ange en siffra 0–99 för {label}."
                    if not (0 <= val <= 99):
                        return None, f"Ange en siffra 0–99 för {label}."
                    return val, ""

                logical_val, err_log = _validate_0_99("logisk färdighet", logical_raw)
                verbal_val, err_verb = _validate_0_99("verbal färdighet", verbal_raw)

                if err_log:
                    context["error"] = err_log
                elif err_verb:
                    context["error"] = err_verb

                context["logical_score"] = logical_val if logical_val is not None else ""
                context["verbal_score"] = verbal_val if verbal_val is not None else ""

                # FIX: spara score i session EFTER validering
                request.session["logical_score"] = context["logical_score"]
                request.session["verbal_score"] = context["verbal_score"]

                # Excel
                if "excel" in request.FILES:
                    try:
                        file = request.FILES["excel"]
                        wb = openpyxl.load_workbook(file)
                        ws = wb.active

                        out = io.StringIO()
                        for row in ws.iter_rows(values_only=True):
                            out.write("\t".join([str(c) if c is not None else "" for c in row]) + "\n")
                        excel_text = out.getvalue()
                    except Exception as e:
                        context["error"] = "Kunde inte läsa excelfilen: " + str(e)[:400]
                else:
                    context["error"] = "Ladda upp en Excelfil."

                # Intervju
                intervju_raw = (request.POST.get("intervju") or "").strip()
                if not intervju_raw and not context["error"]:
                    context["error"] = "Klistra in intervjuanteckningar."

                # CV-text
                cv_raw = (request.POST.get("cv_text") or "").strip()
                if not cv_raw:
                    if not context["error"]:
                        context["error"] = "Klistra in kandidatens CV som text."
                else:
                    # ✅ Spara EXAKT den text användaren klistrat in
                    context["uploaded_files_markdown"] = cv_raw
                    context["uploaded_files_html"] = markdown(cv_raw)

                if context["error"]:
                    step = 1
                else:
                    context["test_text"] = excel_text
                    context["intervju_text"] = intervju_raw

                    try:
                        ratings, dbg = _ratings_from_worksheet(ws)
                        ratings_json_str = json.dumps(ratings, ensure_ascii=False)
                        context["ratings_json"] = ratings_json_str
                        context["ratings_sidebar"] = _build_sidebar_ratings(ratings)
                    except Exception as e:
                        context["error"] = "Kunde inte tolka betyg från Excel: " + str(e)[:400]
                        step = 1

                    if not context["error"]:
                        uploaded_trimmed = _trim(context["uploaded_files_markdown"])

                        if not context["tq_fardighet_text"]:
                            P = Prompt.objects.get(user=owner, name="tq_fardighet").text
                            context["tq_fardighet_text"] = _run_openai(
                                P,
                                style,
                                excel_text=_trim(excel_text),
                                intervju_text=_trim(intervju_raw),
                                ratings_json=ratings_json_str,
                                betygsskala_forklaring=betygsskala_prompt,
                                uploaded_files=uploaded_trimmed,
                                candidate_name=context["candidate_name"],
                                candidate_role=context["candidate_role"],
                                candidate_first_name=context["candidate_first_name"],
                                candidate_last_name=context["candidate_last_name"],

                                job_ad_text=context.get("job_ad_text", ""),
                                motivation_notes=context.get("motivation_notes", ""),
                                logical_score=str(context.get("logical_score") or ""),
                                verbal_score=str(context.get("verbal_score") or ""),

                                # FIX: skicka valda motivationsfaktorer
                                motivation_factors=selected_motivations_for_ai,
                            )

                            if not context["tq_motivation_text"]:
                                P = Prompt.objects.get(user=owner, name="tq_motivation").text
                                context["tq_motivation_text"] = _run_openai(
                                    P,
                                    style,
                                    excel_text=_trim(excel_text),
                                    intervju_text=_trim(intervju_raw),
                                    ratings_json=ratings_json_str,
                                    betygsskala_forklaring=betygsskala_prompt,
                                    uploaded_files=uploaded_trimmed,
                                    candidate_name=context["candidate_name"],
                                    candidate_role=context["candidate_role"],
                                    candidate_first_name=context["candidate_first_name"],
                                    candidate_last_name=context["candidate_last_name"],

                                    job_ad_text=context.get("job_ad_text", ""),
                                    motivation_notes=context.get("motivation_notes", ""),
                                    logical_score=str(context.get("logical_score") or ""),
                                    verbal_score=str(context.get("verbal_score") or ""),

                                    # FIX: skicka valda motivationsfaktorer
                                    motivation_factors=selected_motivations_for_ai,
                                )

                        step = 2

                        if not report_id:
                            rep = _ensure_report(request, context)
                            report_id = str(rep.id)

            # 2 -> 3
            elif step == 2:
                step = 3

            # 3 -> 4
            elif step == 3:
                if not context["leda_text"]:
                    P = Prompt.objects.get(user=owner, name="leda").text
                    context["leda_text"] = _run_openai(
                        P,
                        style,
                        excel_text=_trim(context["test_text"]),
                        intervju_text=_trim(context["intervju_text"]),
                        ratings_json=ratings_json_str,
                        betygsskala_forklaring=betygsskala_prompt,
                        uploaded_files=_trim(
                            context.get("uploaded_files_markdown")
                            or context.get("uploaded_files_text", "")
                        ),
                        candidate_name=context["candidate_name"],
                        candidate_role=context["candidate_role"],
                        candidate_first_name=context["candidate_first_name"],
                        candidate_last_name=context["candidate_last_name"],
                        motivation_notes=context.get("motivation_notes", ""),
                        logical_score=context.get("logical_score", ""),
                        verbal_score=context.get("verbal_score", ""),

                        # FIX: skicka valda motivationsfaktorer
                        motivation_factors=selected_motivations_for_ai,
                    )
                step = 4

            # 4 -> 5
            elif step == 4:
                if not context["mod_text"]:
                    P = Prompt.objects.get(user=owner, name="mod").text
                    context["mod_text"] = _run_openai(
                        P,
                        style,
                        excel_text=_trim(context["test_text"]),
                        intervju_text=_trim(context["intervju_text"]),
                        ratings_json=ratings_json_str,
                        betygsskala_forklaring=betygsskala_prompt,
                        uploaded_files=_trim(
                            context.get("uploaded_files_markdown")
                            or context.get("uploaded_files_text", "")
                        ),
                        candidate_name=context["candidate_name"],
                        candidate_role=context["candidate_role"],
                        candidate_first_name=context["candidate_first_name"],
                        candidate_last_name=context["candidate_last_name"],
                        motivation_notes=context.get("motivation_notes", ""),
                        logical_score=context.get("logical_score", ""),
                        verbal_score=context.get("verbal_score", ""),

                        # FIX
                        motivation_factors=selected_motivations_for_ai,
                    )
                step = 5

            # 5 -> 6
            elif step == 5:
                if not context["sjalkannedom_text"]:
                    P = Prompt.objects.get(user=owner, name="sjalkannedom").text
                    context["sjalkannedom_text"] = _run_openai(
                        P,
                        style,
                        excel_text=_trim(context["test_text"]),
                        intervju_text=_trim(context["intervju_text"]),
                        ratings_json=ratings_json_str,
                        betygsskala_forklaring=betygsskala_prompt,
                        uploaded_files=_trim(
                            context.get("uploaded_files_markdown")
                            or context.get("uploaded_files_text", "")
                        ),
                        candidate_name=context["candidate_name"],
                        candidate_role=context["candidate_role"],
                        candidate_first_name=context["candidate_first_name"],
                        candidate_last_name=context["candidate_last_name"],
                        motivation_notes=context.get("motivation_notes", ""),
                        logical_score=context.get("logical_score", ""),
                        verbal_score=context.get("verbal_score", ""),

                        # FIX
                        motivation_factors=selected_motivations_for_ai,
                    )
                step = 6

            # 6 -> 7
            elif step == 6:
                if not context["strategi_text"]:
                    P = Prompt.objects.get(user=owner, name="strategi").text
                    context["strategi_text"] = _run_openai(
                        P,
                        style,
                        excel_text=_trim(context["test_text"]),
                        intervju_text=_trim(context["intervju_text"]),
                        ratings_json=ratings_json_str,
                        betygsskala_forklaring=betygsskala_prompt,
                        uploaded_files=_trim(
                            context.get("uploaded_files_markdown")
                            or context.get("uploaded_files_text", "")
                        ),
                        candidate_name=context["candidate_name"],
                        candidate_role=context["candidate_role"],
                        candidate_first_name=context["candidate_first_name"],
                        candidate_last_name=context["candidate_last_name"],
                        motivation_notes=context.get("motivation_notes", ""),
                        logical_score=context.get("logical_score", ""),
                        verbal_score=context.get("verbal_score", ""),

                        # FIX
                        motivation_factors=selected_motivations_for_ai,
                    )
                step = 7

            # 7 -> 8
            elif step == 7:
                if not context["kommunikation_text"]:
                    P = Prompt.objects.get(user=owner, name="kommunikation").text
                    context["kommunikation_text"] = _run_openai(
                        P,
                        style,
                        excel_text=_trim(context["test_text"]),
                        intervju_text=_trim(context["intervju_text"]),
                        ratings_json=ratings_json_str,
                        betygsskala_forklaring=betygsskala_prompt,
                        uploaded_files=_trim(
                            context.get("uploaded_files_markdown")
                            or context.get("uploaded_files_text", "")
                        ),
                        candidate_name=context["candidate_name"],
                        candidate_role=context["candidate_role"],
                        candidate_first_name=context["candidate_first_name"],
                        candidate_last_name=context["candidate_last_name"],
                        motivation_notes=context.get("motivation_notes", ""),
                        logical_score=context.get("logical_score", ""),
                        verbal_score=context.get("verbal_score", ""),

                        # FIX
                        motivation_factors=selected_motivations_for_ai,
                    )
                step = 8

            # 8 -> 9
            elif step == 8:
                if not context["sur_text"]:
                    P = Prompt.objects.get(user=owner, name="styrkor_utveckling_risk").text
                    context["sur_text"] = _run_openai(
                        P,
                        style,
                        tq_fardighet_text=context["tq_fardighet_text"],
                        tq_motivation_text=context["tq_motivation_text"],
                        leda_text=context["leda_text"],
                        mod_text=context["mod_text"],
                        sjalkannedom_text=context["sjalkannedom_text"],
                        strategi_text=context["strategi_text"],
                        kommunikation_text=context["kommunikation_text"],
                        uploaded_files=_trim(
                            context.get("uploaded_files_markdown")
                            or context.get("uploaded_files_text", "")
                        ),
                        candidate_name=context["candidate_name"],
                        candidate_role=context["candidate_role"],
                        candidate_first_name=context["candidate_first_name"],
                        candidate_last_name=context["candidate_last_name"],
                        motivation_notes=context.get("motivation_notes", ""),
                        logical_score=context.get("logical_score", ""),
                        verbal_score=context.get("verbal_score", ""),

                        # FIX
                        motivation_factors=selected_motivations_for_ai,
                    )
                step = 9

            # 9 -> 10
            elif step == 9:
                # FIX: loaded_data finns nu alltid definierad (laddades tidigt)
                if loaded_data:
                    context = _apply_report_data_to_context(context, loaded_data)

                if request.method == "POST":
                    pass

                if not context["slutsats_text"]:
                    P = Prompt.objects.get(user=owner, name="sammanfattande_slutsats").text
                    context["slutsats_text"] = _run_openai(
                        P,
                        style,
                        sur_text=context["sur_text"],
                        tq_fardighet_text=context["tq_fardighet_text"],
                        tq_motivation_text=context["tq_motivation_text"],
                        leda_text=context["leda_text"],
                        mod_text=context["mod_text"],
                        sjalkannedom_text=context["sjalkannedom_text"],
                        strategi_text=context["strategi_text"],
                        kommunikation_text=context["kommunikation_text"],
                        uploaded_files=_trim(
                            context.get("uploaded_files_markdown")
                            or context.get("uploaded_files_text", "")
                        ),
                        candidate_name=context["candidate_name"],
                        candidate_role=context["candidate_role"],
                        candidate_first_name=context["candidate_first_name"],
                        candidate_last_name=context["candidate_last_name"],
                        motivation_notes=context.get("motivation_notes", ""),
                        logical_score=context.get("logical_score", ""),
                        verbal_score=context.get("verbal_score", ""),

                        # FIX
                        motivation_factors=selected_motivations_for_ai,
                    )
                step = 10

            # 10 -> 11
            elif step == 10:
                step = 11

    # uppdatera step i context efter POST-logik
    context["step"] = step

    # 🔹 Bygg lista med fulla objekt för de valda motivationsfaktorerna
    selected_motivation_keys = context.get("selected_motivation_keys") or \
                              request.session.get("selected_motivation_keys", [])

    context["selected_motivation_keys"] = selected_motivation_keys
    context["selected_motivations"] = [
        {
            "key": k,
            "label": MOTIVATION_FACTORS[k]["label"],
            "definition": MOTIVATION_FACTORS[k]["definition"],
        }
        for k in selected_motivation_keys
        if k in MOTIVATION_FACTORS
    ]

    # ---------- 4.5) Skapa sidopanelens chattsession (per steg) ----------
    # FIX: undvik att skapa ny session vid varje refresh på steg 1.
    sidebar_key = f"sidebar_session_domarnamnden_{report_id or 'new'}_{context['step']}"
    sidebar_session_id = request.session.get(sidebar_key)

    sidebar_session = None
    if sidebar_session_id:
        sidebar_session = ChatSession.objects.filter(id=sidebar_session_id, user=request.user).first()

    if not sidebar_session:
        # Behåll din logik: steg 1 får en "ny rapport"-titel annars steg X.
        if context["step"] == 1:
            sidebar_session = ChatSession.objects.create(
                user=request.user,
                flow="domarnamnden",
                step=1,
                title="Domarnämnden – ny rapport",
                system_prompt="""
    Du är en skrivassistent som hjälper användaren att förbättra korta textavsnitt
    i ett rapportverktyg.
    """,
            )
        else:
            sidebar_session, _ = ChatSession.objects.get_or_create(
                user=request.user,
                flow="domarnamnden",
                step=context["step"],
                defaults={
                    "title": f"Domarnämnden – steg {context['step']}",
                    "system_prompt": """
    Du är en skrivassistent som hjälper användaren att förbättra korta textavsnitt
    i ett rapportverktyg.
    """,
                },
            )

        request.session[sidebar_key] = sidebar_session.id

    sidebar_messages = ChatMessage.objects.filter(
        session=sidebar_session
    ).order_by("created_at")

    context["sidebar_session_id"] = sidebar_session.id
    context["sidebar_messages"] = sidebar_messages

    # Bygg context som skickas till AI i sidopanelen
    ratings_json_for_sidebar = context.get("ratings_json", ratings_json_str or "")
    sidebar_ctx = _build_sidebar_context(
        owner=owner,
        step=context["step"],
        context=context,
        ratings_json_str=ratings_json_for_sidebar,
    )
    context["sidebar_context_json"] = json.dumps(sidebar_ctx, ensure_ascii=False)

    # ---------- 5) Förbered HTML-versioner till sammanställningen ----------
    context["tq_fardighet_html"]   = _markdown_to_html(context.get("tq_fardighet_text", ""))
    context["sur_html"]            = _markdown_to_html(context.get("sur_text", ""))
    context["tq_motivation_html"]  = _markdown_to_html(context.get("tq_motivation_text", ""))
    context["leda_html"]           = _markdown_to_html(context.get("leda_text", ""))
    context["mod_html"]            = _markdown_to_html(context.get("mod_text", ""))
    context["sjalkannedom_html"]   = _markdown_to_html(context.get("sjalkannedom_text", ""))
    context["strategi_html"]       = _markdown_to_html(context.get("strategi_text", ""))
    context["kommunikation_html"]  = _markdown_to_html(context.get("kommunikation_text", ""))
    context["slutsats_html"]       = _markdown_to_html(context.get("slutsats_text", ""))

    # ---------- 6) Render ----------
    # Alla faktorer till formuläret (input-rutorna)
    context["motivation_factors"] = MOTIVATION_FACTORS

    # Endast valda faktorer till sidopanelen
    selected_motivation_keys = context.get("selected_motivation_keys") or \
                              request.session.get("selected_motivation_keys", [])

    context["selected_motivation_keys"] = selected_motivation_keys
    context["selected_motivations"] = [
        {
            "key": k,
            "label": MOTIVATION_FACTORS[k]["label"],
            "definition": MOTIVATION_FACTORS[k]["definition"],
        }
        for k in selected_motivation_keys
        if k in MOTIVATION_FACTORS
    ]

    # ── Har vi någon data till sidopanelen? ─────────────────────────────
    has_sidebar_data = any([
        context.get("candidate_name"),
        context.get("candidate_role"),
        context.get("logical_score") not in ("", None),
        context.get("verbal_score") not in ("", None),
        context.get("selected_motivations"),
        context.get("ratings_sidebar"),
        context.get("cv_text"),
        context.get("job_ad_text"),
        context.get("motivation_notes"),
        context.get("intervju_text"),
    ])
    context["has_sidebar_data"] = has_sidebar_data

    rep = None

    # Om vi laddade en report via report_id
    if loaded_report:
        rep = loaded_report

    # På POST vill vi alltid spara
    if request.method == "POST" and report_id:
        if not rep:
            rep = _get_report_or_404(report_id)

        _save_report_state(rep, context)

    # Skicka med report_id till templaten så den kan POST:as vidare
    context["report_id"] = report_id

    return render(request, "index.html", context)




# ====== CHAT HELPERS =========================================================
import chardet
from docx import Document
from PyPDF2 import PdfReader

MAX_FILE_TEXT = 15000  # tecken; vi trimmar för att inte spränga tokens

def _read_file_text(django_file) -> str:
    name = django_file.name.lower()
    try:
        # se till att vi läser från början
        if hasattr(django_file, "open"):
            django_file.open(mode="rb")
        try:
            django_file.seek(0)
        except Exception:
            pass

        if name.endswith(".pdf"):
            reader = PdfReader(django_file)
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        elif name.endswith(".docx"):
            doc = Document(django_file)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif name.endswith((".txt", ".csv", ".md", ".json", ".py", ".html")):
            data = django_file.read()
            enc = chardet.detect(data).get("encoding") or "utf-8"
            text = data.decode(enc, errors="ignore")
        else:
            text = ""

        # CRUCIAL för Postgres: inga NUL-tecken i TextField
        return (text or "").replace("\x00", "")
    except Exception:
        return ""

def _trim_middle(s: str, max_chars: int = MAX_FILE_TEXT) -> str:
    s = s or ""
    if len(s) <= max_chars: 
        return s
    return s[: max_chars//2] + "\n...\n" + s[- max_chars//2 :]

def _build_openai_messages(session):
    """Konstruera historiken som OpenAI-meddelanden."""
    msgs = [{"role":"system","content": session.system_prompt}]
    for m in session.messages.order_by("created_at"):
        msgs.append({"role": m.role, "content": m.content})
    return msgs

# ====== CHAT VIEWS ===========================================================
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, redirect
from django.core.files.base import ContentFile
from .models import ChatSession, ChatMessage, ChatAttachment

@login_required
def chat_home(request):
    # skapa session om none
    if request.method == "POST":
        title = request.POST.get("title","Ny chatt").strip() or "Ny chatt"
        s = ChatSession.objects.create(user=request.user, title=title)
        # valfri första system-prompt override
        sp = request.POST.get("system_prompt")
        if sp:
            s.system_prompt = sp
            s.save()
        return redirect("chat_session", session_id=s.id)

    sessions = ChatSession.objects.filter(user=request.user).order_by("-updated_at")
    return render(request, "chat_home.html", {"sessions": sessions})


@login_required
def chat_session(request, session_id):
    session = get_object_or_404(ChatSession, id=session_id, user=request.user)

    # uppdatera system-prompt/titel
    if request.method == "POST" and "save_settings" in request.POST:
        session.title = request.POST.get("title") or session.title
        session.system_prompt = request.POST.get("system_prompt") or session.system_prompt
        session.save()
        return redirect("chat_session", session_id=session.id)

    # skicka meddelande
    if request.method == "POST" and "send_message" in request.POST:
        user_text = request.POST.get("message", "").strip()
        if user_text or request.FILES:
            # 1) Spara själva user-meddelandet som bara "ren" text (utan utdrag)
            user_msg = ChatMessage.objects.create(session=session, role="user", content=user_text)

            # 2) Spara bilagor + extrahera text (men skriv inte in den i user_msg.content)
            file_texts = []
            for f in request.FILES.getlist("files"):
                att = ChatAttachment(message=user_msg, original_name=f.name)
                att.file.save(f.name, f, save=False)  # spara originalfilen
                att_text = _read_file_text(att.file)  # extrahera text
                att.text_excerpt = _trim_middle(att_text, MAX_FILE_TEXT)
                att.save()

                # bygg endast ett "osynligt" prompt-tillägg till AI (inte till UI)
                if att.text_excerpt:
                    file_texts.append(f"\n--- \nFIL: {att.original_name}\n{att.text_excerpt}")

            # 3) Bygg prompt till OpenAI: ersätt sista user-meddelandet med en version
            #    som inkluderar filtexter, men utan att ändra vad som sparas i DB/UI
            try:
                messages = _build_openai_messages(session)
                combined = user_msg.content
                if file_texts:
                    combined += "\n\n(Bifogade filer – textutdrag, visas ej för användaren)" + "".join(file_texts)
                messages[-1]["content"] = combined  # endast för anropet, ej sparat i DB

                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=messages,
                    max_tokens=1200,
                    temperature=0.3,
                )
                ai_text = resp.choices[0].message.content.strip()
            except Exception as e:
                ai_text = f"(Ett fel inträffade vid AI-anropet: {e})"

            # 4) Spara AI-svaret och bumpa sessionen
            ChatMessage.objects.create(session=session, role="assistant", content=ai_text)
            session.save()
            return redirect("chat_session", session_id=session.id)

    messages = session.messages.order_by("created_at")
    sessions = ChatSession.objects.filter(user=request.user).order_by("-updated_at")[:20]
    return render(
        request,
        "chat_session.html",
        {"session": session, "messages": messages, "sessions": sessions}
    )

def _build_sidebar_context_message(ctx: dict) -> str:
    """
    Gör om sidebar_context-json till en text som skickas som extra system-meddelande.
    """
    parts = []
    parts.append(
        "Du är en sidopanel-assistent i ett rapportverktyg för Domarnämnden. "
        "Du hjälper användaren att förbättra, korta ner eller förtydliga texter."
    )

    name = ctx.get("candidate_name") or ""
    role = ctx.get("candidate_role") or ""
    if name or role:
        parts.append(f"Kandidat: {name or 'okänd'}, roll/tjänst: {role or 'ej angiven'}.")

    step = ctx.get("step")
    if step:
        parts.append(f"Aktuellt steg i verktyget: steg {step}.")

    test_text = ctx.get("test_text") or ""
    intervju_text = ctx.get("intervju_text") or ""
    cv_text = ctx.get("cv_text") or ""

    if test_text:
        parts.append("Testdata (kort utdrag):\n" + _trim(test_text, 1200))
    if intervju_text:
        parts.append("Intervjuanteckningar (kort utdrag):\n" + _trim(intervju_text, 1200))
    if cv_text:
        parts.append("CV-text (kort utdrag):\n" + _trim(cv_text, 1200))

    for sec in ctx.get("sections", []):
        label = sec.get("field_label") or sec.get("field_key")
        prompt_name = sec.get("prompt_name")
        prompt_text = sec.get("prompt_text") or ""
        current_text = sec.get("current_text") or ""

        parts.append(
            f"Sektion: {label}.\n"
            f"Prompt som används för denna sektion (namn: {prompt_name}):\n"
            f"{_trim(prompt_text, 1200)}\n\n"
            f"Nuvarande text i denna sektion:\n{_trim(current_text, 1200)}"
        )

    parts.append(
        "När användaren ber dig ändra, förbättra eller korta 'texten', ska du utgå från "
        "den nuvarande texten i den aktuella sektionen och svara med ett färdigt förbättrat "
        "textförslag på svenska, utan extra förklaringar."
    )

    return "\n\n".join(parts)

@csrf_exempt
@login_required
def chat_send(request, session_id):
    """
    POST: message + files  -> streamar AI-svaret som text/plain.
    Kompatibel med OpenAI Python SDK där man använder create(..., stream=True).
    """
    session = get_object_or_404(ChatSession, id=session_id, user=request.user)
    if request.method != "POST":
        return StreamingHttpResponse(
            iter(["Only POST allowed"]),
            content_type="text/plain; charset=utf-8",
            status=405
        )

    user_text = (request.POST.get("message") or "").strip()
    if not user_text and not request.FILES:
        resp = StreamingHttpResponse(
            iter([""]),
            content_type="text/plain; charset=utf-8",
            status=200
        )
        resp["Cache-Control"] = "no-cache"
        resp["X-Accel-Buffering"] = "no"
        return resp

    # 1) Spara user-meddelande (utan filutdrag i content)
    user_msg = ChatMessage.objects.create(session=session, role="user", content=user_text)

    # 2) Filer -> extrahera text -> endast för prompten
    file_texts = []
    for f in request.FILES.getlist("files"):
        att = ChatAttachment(message=user_msg, original_name=f.name)
        att.file.save(f.name, f, save=False)
        att_text = _read_file_text(att.file)
        att.text_excerpt = _trim_middle(att_text, MAX_FILE_TEXT)
        att.save()
        if att.text_excerpt:
            file_texts.append(f"\n--- \nFIL: {att.original_name}\n{att.text_excerpt}")

    # 3) Historik + ersätt sista user content med dold filtext (endast i prompten)
    messages = _build_openai_messages(session)

    # 🧠 Lägg till extra context från verktyget om det skickats med
    sidebar_ctx_raw = request.POST.get("sidebar_context")
    if sidebar_ctx_raw:
        try:
            sidebar_ctx = json.loads(sidebar_ctx_raw)
        except Exception:
            sidebar_ctx = None

        if sidebar_ctx:
            ctx_text = _build_sidebar_context_message(sidebar_ctx)
            # Lägg in efter grund-systemprompten så att den alltid finns med
            messages.insert(1, {"role": "system", "content": ctx_text})

    combined = user_msg.content
    if file_texts:
        combined += "\n\n(Bifogade filer – textutdrag, visas ej för användaren)" + "".join(file_texts)
    messages[-1]["content"] = combined

    def token_stream():
        pieces = []

        try:
            stream = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=messages,
                temperature=0.3,
                max_tokens=1200,
                stream=True,
            )
            for chunk in stream:
                piece = ""
                try:
                    delta = chunk.choices[0].delta
                    if isinstance(delta, dict):
                        piece = delta.get("content") or ""
                    else:
                        piece = getattr(delta, "content", "") or ""
                except Exception:
                    piece = ""

                if piece:
                    pieces.append(piece)
                    yield piece

        except Exception as e:
            err = f"\n\n(Ett fel inträffade vid AI-anropet: {e})"
            pieces.append(err)
            yield err
        finally:
            ai_text = "".join(pieces).strip()
            ChatMessage.objects.create(session=session, role="assistant", content=ai_text)
            session.save()

    resp = StreamingHttpResponse(token_stream(), content_type="text/plain; charset=utf-8")
    resp["Cache-Control"] = "no-cache"
    resp["X-Accel-Buffering"] = "no"
    return resp


@require_POST
@login_required
@csrf_exempt
def sidebar_chat(request):
    """
    Enkel chatt-endpoint för sidopanelen.
    Tar emot:
      - session_id
      - message
      - context  (texten i nuvarande steg / textarea)
    Returnerar JSON: {"reply": "..."}
    """
    session_id = request.POST.get("session_id")
    message = (request.POST.get("message") or "").strip()
    context_blob = (request.POST.get("context") or "").strip()

    if not session_id or not message:
        return JsonResponse({"error": "session_id och message krävs"}, status=400)

    session = get_object_or_404(ChatSession, id=session_id, user=request.user)

    # 1) Spara användarens meddelande
    user_msg = ChatMessage.objects.create(session=session, role="user", content=message)

    # 2) Bygg historik för OpenAI
    messages = _build_openai_messages(session)

    # Lägg in kontexten snyggt innan användarens text
    combined = message
    if context_blob:
        combined = (
            "Du är en assistent som hjälper användaren att förbättra texten i ett verktyg.\n"
            "Nedanför finns den text som hör till det steg användaren jobbar i just nu.\n"
            "När användaren skriver saker som 'den här texten', 'det här stycket' etc, "
            "så syftar de på texten nedan.\n\n"
            "=== AKTUELL TEXT I VERKTYGET ===\n"
            f"{context_blob}\n\n"
            "=== ANVÄNDARENS MEDDELANDE ===\n"
            f"{message}"
        )

    messages.append({"role": "user", "content": combined})

    # 3) Anropa OpenAI
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0.3,
            max_tokens=800,
        )
        reply = resp.choices[0].message.content.strip()
    except Exception as e:
        reply = f"(Ett fel inträffade i chatten: {e})"

    # 4) Spara AI-svaret
    ChatMessage.objects.create(session=session, role="assistant", content=reply)
    session.save()

    return JsonResponse({"reply": reply})

@require_POST
@login_required
def chat_delete(request, session_id):
    s = get_object_or_404(ChatSession, id=session_id, user=request.user)
    s.delete()  # Messages/attachments följer med om du har on_delete=CASCADE
    return redirect("chat_home")


@login_required
def report_list(request):
    reports = Report.objects.filter(deleted_at__isnull=True).order_by("-updated_at")
    return render(request, "report_list.html", {"reports": reports})


@login_required
def report_open(request, report_id):
    rep = _get_report_or_404(report_id)
    data = rep.data or {}
    # visa "slutresultat" (du kan göra en proper template senare)
    # här skickar vi samma keys som du redan renderar i index.html sammanställning
    context = {
        "report": rep,
        "data": data,
        "tq_fardighet_html": _markdown_to_html(data.get("tq_fardighet_text", "")),
        "tq_motivation_html": _markdown_to_html(data.get("tq_motivation_text", "")),
        "leda_html": _markdown_to_html(data.get("leda_text", "")),
        "mod_html": _markdown_to_html(data.get("mod_text", "")),
        "sjalkannedom_html": _markdown_to_html(data.get("sjalkannedom_text", "")),
        "strategi_html": _markdown_to_html(data.get("strategi_text", "")),
        "kommunikation_html": _markdown_to_html(data.get("kommunikation_text", "")),
        "sur_html": _markdown_to_html(data.get("sur_text", "")),
        "slutsats_html": _markdown_to_html(data.get("slutsats_text", "")),
    }
    return render(request, "report_open.html", context)


@login_required
def report_edit(request, report_id):
    report = get_object_or_404(Report, id=report_id, created_by=request.user)

    # Anta att report.data är en dict (JSONField) med ALLT du behöver
    # Ex: kandidat, intervju_text, excel_text, motivation_factors, etc.
    report_data = report.data or {}

    # Lägg wizard-state i sessionen
    request.session["wizard_mode"] = "edit"
    request.session["wizard_report_id"] = str(report.id)
    request.session["wizard_step"] = 2
    request.session["wizard_data"] = report_data

    # Viktigt: markera att sessionen ändrats
    request.session.modified = True

    return redirect(f"{reverse('index')}?report_id={report.id}&step=2")


@login_required
@require_POST
def report_delete(request, report_id):
    rep = _get_report_or_404(report_id)
    rep.deleted_at = timezone.now()
    rep.save(update_fields=["deleted_at"])
    return redirect("report_list")


@login_required
def report_download(request, report_id):
    """
    Återanvänd din Word-export, fast från sparad data.
    Vi simulerar POST-build_doc med rep.data.
    """
    rep = _get_report_or_404(report_id)
    data = rep.data or {}

    # Bygg ett 'fake context' som matchar din build_doc-del
    context = data.copy()
    context["candidate_name"] = context.get("candidate_name", "")  # safety

    # samma kod som i build_doc men med data istället för request.POST:
    from django.http import HttpResponse
    from docx import Document

    template_path = os.path.join(settings.BASE_DIR, "reports", "domarnamnden_template.docx")
    doc = Document(template_path)

    # om du vill spara bilder senare kan du lägga dem i rep.data också
    leda_image_data = data.get("leda_image", "")
    mod_image_data = data.get("mod_image", "")
    sjalkannedom_image_data = data.get("sjalkannedom_image", "")
    strategi_image_data = data.get("strategi_image", "")
    kommunikation_image_data = data.get("kommunikation_image", "")

    # motivations-text (som du redan bygger)
    selected_motivation_keys = data.get("selected_motivation_keys", []) or []
    motivation_lines = []
    for k in selected_motivation_keys:
        m = MOTIVATION_FACTORS.get(k)
        if not m:
            continue
        motivation_lines.append(f"{m['label']}\n{m['definition']}".strip())
    selected_motivations_text = "\n\n".join(motivation_lines)

    mapping = {
        "{candidate_name}": context.get("candidate_name", ""),
        "{candidate_first_name}": context.get("candidate_first_name", ""),
        "{candidate_last_name}": context.get("candidate_last_name", ""),
        "{candidate_role}": context.get("candidate_role", ""),
        "{tq_fardighet_text}": html_to_text(context.get("tq_fardighet_text", "")),
        "{sur_text}": html_to_text(context.get("sur_text", "")),
        "{tq_motivation_text}": html_to_text(context.get("tq_motivation_text", "")),
        "{leda_text}": html_to_text(context.get("leda_text", "")),
        "{mod_text}": html_to_text(context.get("mod_text", "")),
        "{sjalkannedom_text}": html_to_text(context.get("sjalkannedom_text", "")),
        "{strategi_text}": html_to_text(context.get("strategi_text", "")),
        "{kommunikation_text}": html_to_text(context.get("kommunikation_text", "")),
        "{selected_motivations}": selected_motivations_text,
    }
    docx_replace_text(doc, mapping)

    ratings_json_raw = data.get("ratings_json") or ""
    ratings_doc = {}
    if isinstance(ratings_json_raw, dict):
        ratings_doc = ratings_json_raw
    elif isinstance(ratings_json_raw, str) and ratings_json_raw.strip():
        try:
            ratings_doc = json.loads(ratings_json_raw)
        except Exception:
            ratings_doc = {}

    if ratings_doc:
        replace_table_placeholder(doc, "{leda_table}", ratings_doc, "leda_utveckla_och_engagera")
        replace_table_placeholder(doc, "{mod_table}", ratings_doc, "mod_och_handlingskraft")
        replace_table_placeholder(doc, "{sjalkannedom_table}", ratings_doc, "sjalkannedom_och_emotionell_stabilitet")
        replace_table_placeholder(doc, "{strategi_table}", ratings_doc, "strategiskt_tankande_och_anpassningsformaga")
        replace_table_placeholder(doc, "{kommunikation_table}", ratings_doc, "kommunikation_och_samarbete")

    replace_image_placeholder(doc, "{leda_image}", leda_image_data)
    replace_image_placeholder(doc, "{mod_image}", mod_image_data)
    replace_image_placeholder(doc, "{sjalkannedom_image}", sjalkannedom_image_data)
    replace_image_placeholder(doc, "{strategi_image}", strategi_image_data)
    replace_image_placeholder(doc, "{kommunikation_image}", kommunikation_image_data)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    filename = f"bedomning_{context.get('candidate_name','rapport')}.docx"
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    doc.save(response)
    return response
